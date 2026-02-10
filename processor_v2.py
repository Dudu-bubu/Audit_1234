"""
GDP Document Reviewer - Business Logic
Contains DocumentProcessor and Validator classes with all validation rules.
"""

import os
import re
import json
import time
import uuid
import logging
import requests
from datetime import datetime
from spellchecker import SpellChecker
from typing import List, Dict, Any, Optional, Set, Tuple, Annotated
from concurrent.futures import ThreadPoolExecutor, as_completed
from dataclasses import dataclass, field
from json import JSONDecodeError
from collections import defaultdict
import dateutil.parser as date_parser

import docx
from docx.oxml.text.paragraph import CT_P
from docx.oxml.table import CT_Tbl
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX
from openai import AzureOpenAI
from pydantic import BaseModel, Field
from difflib import SequenceMatcher

from models import (
    DocumentModel, SectionModel, ParagraphModel, TableModel, Finding, ValidationResult,
    AcronymManager, GrammarChecker, PromptBuilder, get_priority, get_priority_color,
    add_comment_to_paragraph, heading_similarity, normalize_heading_text,
    GUIDELINE_MATRIX, GUIDELINE_CATEGORY_MAP
)

logger = logging.getLogger(__name__)

# ============================================================================
# DOCUMENT PROCESSOR
# ============================================================================
print("NUCLEAR_VERSION_TAG: V_FINAL_DEDUPE_LAYER")

def clean_finding_text(text: str) -> str:
    """Bulletproof cleaner that reduces any finding to its raw substrate by peeling all labels."""
    if not text: return ""
    import re
    
    # 0. FILE LOGGING DEBUG (Temporary)
    try:
        with open("cleaner_debug.log", "a", encoding="utf-8") as debug_f:
             debug_f.write(f"\n--- CLEAN START: {text[:50]}... ---\n")
    except: pass
    
    # 1. CHARACTER-LEVEL CLEANING
    raw = str(text).strip()
    # Remove BOMs and hidden unicode debris
    for b in ['\ufeff', '\u200b', '\u200c', '\u200d', '\u200e', '\u200f']:
        raw = raw.replace(b, '')
    
    # 2. ITERATIVE LITERAL & REGEX PEELING
    # We do literal first because it's safer, then regex for fuzzy matches.
    current = raw
    for _ in range(8): # Peel up to 8 layers
        prev = current
        
        # A. Remove exact literal phrases we are known to use
        low_curr = current.lower()
        
        # Strip exact common prefixes (case insensitive literals)
        literals = [
            "observation:", "recommendation:", "issue:", "suggestion:", 
            "sugg:", "obs:", "issue identified by ai:", "issue found:"
        ]
        for lit in literals:
            if current.lower().startswith(lit):
                current = current[len(lit):].strip()
        
        # B. Strip "Priority" brackets: [Priority is X], (Priority is X), etc.
        # Handles multi-line or weird spacing, and swallows trailing separators like . - :
        # REMOVED ANCHOR ^ to catch it even if junk precedes it
        current = re.sub(r'\s*[\[\(][^\]\)]*?Priority[^\]\)]*?[\]\)]\s*[\:\-\.]?\s*', '', current, count=1, flags=re.IGNORECASE | re.DOTALL).strip()
        current = re.sub(r'\s*\[\s*(BLOCKING|HIGH|MEDIUM|LOW)\s*\]\s*[\:\-\.]?\s*', '', current, count=1, flags=re.IGNORECASE).strip()
        
        # C. Strip Structural Markers: [Row X, Col Y], [Section X]
        current = re.sub(r'^\s*\[\s*(Row|Col|Section|Table)\s+\d+.*?\s*\]\s*', '', current, flags=re.IGNORECASE).strip()
        
        # D. Generic Label Regex (Fuzzy)
        # REMOVED ANCHOR ^ to catch "   Observation:"
        current = re.sub(r'\s*(Observation|Recommendation|OBS|Issue|Suggestion|Sugg)\s*[:\-]\s*', '', current, count=1, flags=re.IGNORECASE).strip()
        
        if current == prev: 
            # EMERGENCY BREAK: Check if we still start with "Observation" despite regex failure
            # CASE INSENSITIVE SPLIT
            if re.match(r'^\s*observation\s*:', current, re.IGNORECASE):
                # Split using regex to handle case variance
                current = re.split(r'observation\s*:', current, maxsplit=1, flags=re.IGNORECASE)[-1].strip()
            elif current.strip().lower().startswith("[priority"):
                 # Force peel bracket
                 if "]" in current: current = current.split("]", 1)[-1].strip()
            else:
                break
        
    # 3. FINAL CLEANUP
    # Remove leading separators like " | " or trailing artifacts after stripping
    current = re.sub(r'^[:\-\s\|]+', '', current).strip()

    try:
        with open("cleaner_debug.log", "a", encoding="utf-8") as debug_f:
             if current != text.strip():
                 debug_f.write(f"CHANGED: '{text[:30]}...' -> '{current[:30]}...'\n")
             else:
                 debug_f.write(f"NO_CHANGE: '{current[:30]}...'\n")
    except: pass
        
    return current

class DocumentProcessor:
    """Handles all document processing operations"""
    
    def __init__(self):
        pass

    def extract_font_properties(self, para: docx.text.paragraph.Paragraph) -> Tuple[str, float, bool]:
        """Extract font name, size, and bold status from a paragraph."""
        font_name = "Unknown"
        font_size = 0.0
        is_bold = False
        
        # Check all runs to find the dominant font properties
        # Priority: First run with explicit formatting, then style-based
        if para.runs:
            for run in para.runs:
                if not run.text.strip():
                    continue
                    
                if run.font.name:
                    font_name = run.font.name
                else:
                    # Check Run XML for Theme Font
                    try:
                        rPr = run._element.rPr
                        if rPr is not None and rPr.rFonts is not None:
                            xml_str = str(rPr.rFonts.xml)
                            if 'asciiTheme' in xml_str or 'hAnsiTheme' in xml_str:
                                if 'major' in xml_str.lower():
                                    font_name = "Word Theme Font (Major/Calibri Light)"
                                else:
                                    font_name = "Word Theme Font (Minor/Calibri)"
                    except:
                        pass
                    
                if run.font.size:
                    font_size = run.font.size.pt
                if not is_bold and run.font.bold:
                    is_bold = True
                
                if font_name != "Unknown" and font_size > 0:
                    break
        
        if (font_name == "Unknown" or font_size == 0):
            # Check paragraph element properties (pPr -> rPr)
            try:
                rPr = para._element.pPr.rPr
                if rPr is not None and rPr.rFonts is not None:
                     xml_str = str(rPr.rFonts.xml)
                     if 'asciiTheme' in xml_str or 'hAnsiTheme' in xml_str:
                         if 'major' in xml_str.lower():
                             font_name = "Word Theme Font (Major/Calibri Light)"
                         else:
                             font_name = "Word Theme Font (Minor/Calibri)"
            except:
                pass

        if (font_name == "Unknown" or font_size == 0):
            style = para.style
            while style and (font_name == "Unknown" or font_size == 0):
                try:
                    style_font = style.font
                    if font_name == "Unknown":
                        if style_font.name:
                            font_name = style_font.name
                        else:
                            try:
                                rPr = style.element.rPr
                                if rPr is not None and rPr.rFonts is not None:
                                    xml_str = str(rPr.rFonts.xml)
                                    if 'asciiTheme' in xml_str or 'hAnsiTheme' in xml_str:
                                        if 'major' in xml_str.lower():
                                            font_name = "Word Theme Font (Major/Calibri Light)"
                                        else:
                                            font_name = "Word Theme Font (Minor/Calibri)"
                            except:
                                pass

                    if font_size == 0 and style_font.size:
                        font_size = style_font.size.pt
                    if not is_bold and style_font.bold:
                        is_bold = True
                    
                    if font_name != "Unknown":
                        pass 

                    style = style.base_style
                except Exception:
                    break

        return font_name, font_size, is_bold

    def get_paragraph_alignment(self, para):
        alignment = "Left"
        if para.alignment == WD_ALIGN_PARAGRAPH.CENTER: alignment = "Center"
        elif para.alignment == WD_ALIGN_PARAGRAPH.RIGHT: alignment = "Right"
        elif para.alignment == WD_ALIGN_PARAGRAPH.JUSTIFY: alignment = "Justify"
        return alignment

    def extract_hyperlinks(self, paragraph) -> List[str]:
        """Extract hyperlinks from the paragraph."""
        hyperlinks = []
        try:
            iprels = paragraph._element.part.rels
            for rel_id, rel in iprels.items():
                if "hyperlink" in rel.reltype.lower() and rel.target_mode == "External":
                    hyperlinks.append(rel.target_ref)
        except Exception:
            pass
        return hyperlinks

    def detect_heading(self, style_name: str, text: str, para: docx.text.paragraph.Paragraph) -> Tuple[bool, int]:
        """Detect if the paragraph is a heading and return the heading level."""
        is_heading = False
        level = 0
        
        # 1. Check Document Style
        if style_name.startswith('Heading'):
            try:
                level = int(style_name.replace('Heading ', ''))
                is_heading = True
            except ValueError:
                pass
        
        # 2. Heuristics (Formatting and Numbering)
        if not is_heading and text and len(text) < 200:
             font_name, font_size, is_bold = self.extract_font_properties(para)
             
             # Font-based detection (matching GDP standards)
             if font_size >= 23.5: # 24pt
                 is_heading = True
                 level = 1
             elif font_size >= 17.5: # 18pt
                 is_heading = True
                 level = 2
             elif font_size >= 15.5: # 16pt
                 is_heading = True
                 level = 3
             
             # Numbering-based detection (e.g. 1.0, 1.1)
             if not is_heading:
                 if re.match(r'^\d+\.0\b', text):
                     is_heading = True
                     level = 2
                 elif re.match(r'^\d+\.\d+(\.\d+)*\b', text):
                     # 1.1 -> H3, 1.1.1 -> H4
                     dots = text.split()[0].count('.')
                     is_heading = True
                     level = dots + 1
             
             # Fallback fallback
             if not is_heading and is_bold and font_size >= 12.0:
                 is_heading = True
                 level = 1 if font_size >= 14.0 else 2

        return is_heading, level

    def parse_paragraph(self, para: docx.text.paragraph.Paragraph, para_index: int, global_index: int = 0, heading_level: int = 0) -> ParagraphModel:
        """Parse a paragraph and return its model."""
        text = para.text.strip()
        style_name = para.style.name
        font_name, font_size, is_bold = self.extract_font_properties(para)
        alignment = self.get_paragraph_alignment(para)
        
        numbering = None
        try:
            pPr = para._element.pPr
            if pPr is not None and pPr.numPr is not None:
                 numbering = "list"
        except Exception:
            pass
            
        return ParagraphModel(
            id=f"para_{para_index}",
            index=para_index,
            global_index=global_index,
            text=text,
            style=style_name,
            font_name=font_name,
            font_size=font_size or 0.0,
            is_bold=is_bold,
            alignment=alignment,
            numbering=numbering,
            heading_level=heading_level
        )
    
    def check_table_header_repetition(self, table: docx.table.Table) -> bool:
         is_header_repeated = False
         if len(table.rows) > 0:
             try:
                 tr_props = table.rows[0]._tr.trPr
                 if tr_props is not None:
                     header_tag = tr_props.find(docx.oxml.ns.qn('w:tblHeader'))
                     if header_tag is not None:
                         is_header_repeated = True
             except Exception:
                 pass
         return is_header_repeated

    def parse_table(self, table: docx.table.Table, table_index: int, global_index: int = 0) -> Tuple[TableModel, List[ParagraphModel]]:
        """Parse a table and return its model and extracted cell paragraphs."""
        table_id = f"table_{table_index}"
        
        headers = []
        if len(table.rows) > 0:
            for cell in table.rows[0].cells:
                headers.append(cell.text.strip())
        
        table_content = []
        hyperlinks = []
        has_empty = False

        extracted_paragraphs = []

        for r_idx, row in enumerate(table.rows):
            row_data = []
            for c_idx, cell in enumerate(row.cells):
                text = cell.text.replace('\xa0',' ').strip()
                
                # Handle Nested Tables
                if cell.tables:
                    for nested_table in cell.tables:
                        for nr in nested_table.rows:
                            for nc in nr.cells:
                                text += " " + nc.text.replace('\xa0',' ').strip()
                                
                                # Extract hyperlinks from nested table paragraphs
                                for p in nc.paragraphs:
                                     hyperlinks.extend(self.extract_hyperlinks(p))

                row_data.append(text)
                
                # Parse cell paragraphs for individual validation
                # ID format: table_{i}_r{j}_c{k}
                if cell.paragraphs:
                    cell_para = cell.paragraphs[0]
                    p_model = self.parse_paragraph(cell_para, 0) # Index 0 dummy
                    p_model.id = f"{table_id}_r{r_idx}_c{c_idx}"
                    p_model.text = text # Use full cell text
                    extracted_paragraphs.append(p_model)

                # Extract hyperlinks from main cell paragraphs
                for p in cell.paragraphs:
                    hyperlinks.extend(self.extract_hyperlinks(p))
            
            table_content.append(row_data)
            
        # Check empty cells
        for row_data in table_content:
             for cell_text in row_data:
                 if not cell_text:
                     has_empty = True

        # Header repetition
        is_header_repeated = self.check_table_header_repetition(table)

        t_model = TableModel(
            id=table_id,
            index=table_index,
            global_index=global_index,
            row_count=len(table.rows),
            col_count=len(table.columns),
            headers=headers,
            has_empty_cells=has_empty,
            metadata={"is_header_repeated": is_header_repeated, "hyperlinks": list(set(hyperlinks))},
            content=table_content
        )
        return t_model, extracted_paragraphs


    
    def parse_document(self, file_path: str) -> DocumentModel:
        """Parse a Word document and extract structure"""
        logger.info(f"DEBUG: Parsing document: {file_path}")
        try:
            doc = docx.Document(file_path)
        except Exception as e:
            logger.error(f"ERROR: Failed to open docx: {e}")
            raise
        
        logger.info("DEBUG: Docx opened successfully")
        
        model = DocumentModel()
        try:
             model.metadata = {
                "author": doc.core_properties.author,
                "last_modified_by": doc.core_properties.last_modified_by,
                "version": str(doc.core_properties.version),
                "revision_number": str(doc.core_properties.revision),
                "created": doc.core_properties.created.isoformat() if doc.core_properties.created else None,
                "modified": doc.core_properties.modified.isoformat() if doc.core_properties.modified else None
            }
        except Exception:
            pass

        # Root Section
        current_section = SectionModel(id="root", heading_text="Root", heading_level=0, global_index=0)
        model.sections.append(current_section)

        def iter_block_elements(parent):
            """Recursive iterator to find all paragraphs and tables, including those in SDTs."""
            for child in parent.iterchildren():
                if isinstance(child, CT_P):
                    yield ('paragraph', child)
                elif isinstance(child, CT_Tbl):
                    yield ('table', child)
                elif child.tag.endswith('sdt'):
                    content = child.find(qn('w:sdtContent'))
                    if content is not None:
                        yield from iter_block_elements(content)

        para_index = 0
        table_index = 0
        global_counter = 0

        for elem_type, element in iter_block_elements(doc.element.body):
            if elem_type == 'paragraph':
                para = docx.text.paragraph.Paragraph(element, doc)
                text = para.text.strip()
                
                is_heading, level = self.detect_heading(para.style.name, text, para)
                
                if is_heading:
                    new_section = SectionModel(
                        id=str(uuid.uuid4()),
                        heading_text=text,
                        heading_level=level,
                        global_index=global_counter
                    )
                    model.sections.append(new_section)
                    current_section = new_section
                    model.toc_entries.append(text)
                
                p_model = self.parse_paragraph(para, para_index, global_index=global_counter, heading_level=level if is_heading else 0)
                model.paragraphs[p_model.id] = p_model
                current_section.content_ids.append(p_model.id)
                para_index += 1
                global_counter += 1

                # Extract hyperlinks from paragraph relations
                model.hyperlinks.extend(self.extract_hyperlinks(para))
                
            elif elem_type == 'table':
                table = docx.table.Table(element, doc)
                t_model, t_paras = self.parse_table(table, table_index, global_index=global_counter)
                
                # Add extracted cell paragraphs to model for validation
                for tp in t_paras:
                    model.paragraphs[tp.id] = tp
                
                # Aggregate hyperlinks from table metadata to the main model
                if t_model.metadata and "hyperlinks" in t_model.metadata:
                    model.hyperlinks.extend(t_model.metadata["hyperlinks"])
                
                model.tables[t_model.id] = t_model
                current_section.content_ids.append(t_model.id)
                table_index += 1
                global_counter += 1
        
        return model

    def check_structure(self, input_model: DocumentModel, template_model: DocumentModel = None) -> ValidationResult:
        result = ValidationResult()
        findings = []
        
        logger.info(f"DEBUG: check_structure called with {len(input_model.sections)} sections in input model")
        logger.info(f"DEBUG: Section details: {[(s.heading_text, s.heading_level, len(s.content_ids)) for s in input_model.sections]}")
        
        if template_model:
            logger.info("DEBUG: Performing template-driven structural analysis...")
            # Use template as the definitive guide
            score, sim_findings = self._check_toc_similarity(input_model, template_model)
            findings.extend(sim_findings)
            result.similarity_score = score
        else:
            # If no template, we don't assume any sections are mandatory (TOC, References, Acronyms, etc.)
            # as per user instruction: "don't give these sections are amndontary in the input file"
            logger.info("DEBUG: No template provided, skipping structural section checks.")
        
        # General checks that apply regardless of template
        findings.extend(self._check_fonts(input_model))
        findings.extend(self._check_sections(input_model))
        findings.extend(self._check_tables(input_model))
        # findings.extend(self._check_dates_in_document(input_model)) # DUPLICATE: Handled by validate_date_format_and_values
        findings.extend(self._check_whitespace_in_document(input_model))
        
        # NOTE: Prefixes are now added ONLY in Validator.validate_document()
        # Do NOT add prefixes here - it causes duplication!

        result.structure_violations.extend(findings)
        return result
    
    def _check_toc(self, model: DocumentModel) -> Tuple[bool, List[Finding]]:
        findings = []
        toc_section = None
        for section in model.sections:
            if section.heading_text and section.heading_text.lower() in ['table of contents', 'contents', 'toc']:
                toc_section = section
                break
        
        if not toc_section:
            findings.append(Finding(category="Structure", rule_id="MISSING_TOC", issue="Document is missing a Table of Contents", location_id="document", suggestion="Please add a Table of Contents section"))
            return False, findings
        return True, findings
    
    def _check_toc_similarity(self, input_model: DocumentModel, template_model: DocumentModel) -> Tuple[float, List[Finding]]:
        """Compares input document structure against template structure based on headings."""
        findings = []
        if not template_model or not template_model.toc_entries: 
            return 100.0, findings

        # Normalize entries for robust comparison
        input_entries = {normalize_heading_text(e): e for e in input_model.toc_entries if e.strip()}
        template_entries = {normalize_heading_text(e): e for e in template_model.toc_entries if e.strip()}
        
        # Map normalized input headings to their paragraph IDs (for highlighting)
        norm_to_para = {}
        for section in input_model.sections:
            if section.heading_level > 0 and section.content_ids:
                nt = normalize_heading_text(section.heading_text)
                if nt not in norm_to_para:
                    norm_to_para[nt] = section.content_ids[0]

        exact_matches = set(input_entries.keys()) & set(template_entries.keys())
        
        # 1. Analyze INPUT Centric Analysis: Check for Extra Sections
        extra_normalized = set(input_entries.keys()) - exact_matches
        for norm_input in extra_normalized:
            input_orig = input_entries[norm_input]
            loc_id = norm_to_para.get(norm_input, "document")
            
            # Check for fuzzy match
            is_fuzzy_match = False
            best_fuzzy_match = None
            best_fuzzy_score = 0.0
            
            for norm_template in template_entries.keys():
                score = heading_similarity(input_orig, template_entries[norm_template])
                if score >= 0.8:
                    is_fuzzy_match = True
                    if score > best_fuzzy_score:
                        best_fuzzy_score = score
                        best_fuzzy_match = template_entries[norm_template]
                    break
            
            if is_fuzzy_match:
                findings.append(Finding(
                    category="Structure", 
                    rule_id="TOC_FUZZY_MATCH", 
                    issue=f"Section '{input_orig}' identified but differs slightly from template section '{best_fuzzy_match}' ({best_fuzzy_score:.0%}).", 
                    location_id=loc_id, 
                    suggestion=f"Please verify if '{input_orig}' should be aligned with the template heading '{best_fuzzy_match}'."
                ))
            else:
                # 3. EXTRA SESSIONS (In Input but not Template)
                findings.append(Finding(
                    category="Structure", 
                    rule_id="EXTRA_SECTION", 
                    issue=f"Extra section '{input_orig}' found in input document that is not in the template.", 
                    location_id=loc_id, 
                    suggestion="Please remove this section to align with the template, or provide a justification if it is required for this specific document."
                ))

        # 2. TEMPLATE Centric Analysis: Check for Missing Sections
        template_list = [e for e in template_model.toc_entries if e.strip()]
        missing_normalized = set(template_entries.keys()) - exact_matches
        
        for norm_template in missing_normalized:
            template_orig = template_entries[norm_template]
            
            has_fuzzy = False
            for norm_input in extra_normalized:
                if heading_similarity(template_orig, input_entries[norm_input]) >= 0.8:
                    has_fuzzy = True
                    break
            
            if not has_fuzzy:
                location_msg = ""
                loc_id = "document" # Default fallback
                found_context = False
                
                try:
                    t_idx = template_list.index(template_orig)
                    
                    found_context = False
                    loc_id = "document"
                    location_msg = ""
                    
                    # STRATEGY: Find the "Successor" in the input document (the next section that SHOULD follow the missing one).
                    # Then anchor to the section IMMEDIATELY PRECEDING that Successor in the Input.
                    # CRITICAL: Anchor to the last NON-EMPTY paragraph of that predecessor to ensure the comment is visible/clickable.
                    
                    # 1. Find Successor
                    successor_section = None
                    successor_name = ""
                    
                    for i in range(t_idx + 1, len(template_list)):
                        next_t_norm = normalize_heading_text(template_list[i])
                        
                        # Find matching input section
                        match_sec = None
                        if next_t_norm in exact_matches:
                            for s in input_model.sections:
                                if normalize_heading_text(s.heading_text) == next_t_norm:
                                    match_sec = s
                                    break
                        else:
                            for s in input_model.sections:
                                if s.heading_level > 0:
                                    if heading_similarity(template_list[i], s.heading_text) >= 0.8:
                                        match_sec = s
                                        break
                        
                        if match_sec:
                            successor_section = match_sec
                            successor_name = match_sec.heading_text
                            break
                    
                    # 2. Determine Anchor Point based on Successor
                    if successor_section:
                        try:
                            succ_idx = input_model.sections.index(successor_section)
                            
                            if succ_idx > 0:
                                # Anchor to the predecessor section
                                anchor_section = input_model.sections[succ_idx - 1]
                                
                                # Find valid non-empty paragraph in predecessor
                                valid_loc_id = None
                                if anchor_section.content_ids:
                                    # Walk backwards to find text
                                    for pid in reversed(anchor_section.content_ids):
                                        p_obj = input_model.paragraphs.get(pid)
                                        if p_obj and p_obj.text and p_obj.text.strip():
                                            valid_loc_id = pid
                                            break
                                    # Fallback to header if no content text found
                                    if not valid_loc_id:
                                        valid_loc_id = anchor_section.content_ids[0]

                                if valid_loc_id:
                                    loc_id = valid_loc_id
                                    found_context = True
                                    location_msg = f" after the '{anchor_section.heading_text}' section and before '{successor_name}'"
                                    print(f"DEBUG: Anchoring '{template_orig}' to GAP: After '{anchor_section.heading_text}' (Para {loc_id}) -> Before '{successor_name}'")
                            else:
                                # Successor is first. Anchor to its header.
                                if successor_section.content_ids:
                                    loc_id = successor_section.content_ids[0]
                                    found_context = True
                                    location_msg = f" before the '{successor_name}' section"
                                    print(f"DEBUG: Anchoring '{template_orig}' to START: Before '{successor_name}' ({loc_id})")
                        except ValueError:
                            pass
                            
                    else:
                        # 3. No Successor -> End of Document
                        # Find last non-empty paragraph of document
                        if input_model.sections:
                             last_sec = input_model.sections[-1]
                             valid_loc_id = None
                             if last_sec.content_ids:
                                 for pid in reversed(last_sec.content_ids):
                                     p_obj = input_model.paragraphs.get(pid)
                                     if p_obj and p_obj.text and p_obj.text.strip():
                                         valid_loc_id = pid
                                         break
                                 if not valid_loc_id:
                                     valid_loc_id = last_sec.content_ids[0]
                                     
                             if valid_loc_id:
                                 loc_id = valid_loc_id
                                 found_context = True
                                 location_msg = " at the end of the document"
                                 print(f"DEBUG: Anchoring '{template_orig}' to END: After '{last_sec.heading_text}' (Para {loc_id})")

                except Exception as e:
                    logger.error(f"Error finding context for missing section '{template_orig}': {e}")
                
                findings.append(Finding(
                category="Structure", 
                rule_id="MISSING_SECTION", 
                issue=f"Required section '{template_orig}' from the template is missing.", 
                location_id=loc_id, 
                suggestion=f"Please add the missing section '{template_orig}'{location_msg} to align with the template structure."
            ))

        
        if template_entries:
            matching = len(exact_matches)
            similarity = (matching / len(template_entries)) * 100
            if similarity < 50:
                findings.append(Finding(
                    category="Structure", 
                    rule_id="MAJOR_STRUCTURE_MISMATCH", 
                    issue=f"Document structure differs significantly from template ({similarity:.1f}% exact match).", 
                    location_id="document", 
                    suggestion="Review the overall document structure to ensure better alignment with the approved template."
                ))
        
        # Run deeper structural checks if needed
        findings.extend(self._check_heading_structure(input_model, template_model))
        
        # New: Get table stats for composite score
        table_matches, table_total, table_findings = self._check_table_structure(input_model, template_model)
        findings.extend(table_findings)
        
        final_similarity = 100.0
        heading_matches = len(exact_matches)
        heading_total = len(template_entries)
        
        composite_total = heading_total + table_total
        composite_matches = heading_matches + table_matches
        
        if composite_total > 0:
            final_similarity = (composite_matches / composite_total) * 100

        # Adjust LOW_TOC_SIMILARITY check strictly on the final composite score
        # Remove any potential previous duplicate finding if we want to be clean, but simplest is to just rely on this final calc.
        
        return final_similarity, findings

    
    def _check_heading_structure(self, input_model: DocumentModel, template_model: DocumentModel) -> List[Finding]:
        findings = []
        input_sections = [(s.heading_text.lower().strip(), s.heading_level, s.content_ids[0] if s.content_ids else "document") for s in input_model.sections if s.heading_level > 0]
        template_sections = [(s.heading_text.lower().strip(), s.heading_level) for s in template_model.sections if s.heading_level > 0]
        
        template_order = {name: idx for idx, (name, _) in enumerate(template_sections)}
        input_info = {name: (idx, lid) for idx, (name, _, lid) in enumerate(input_sections)}
        common_sections = set(template_order.keys()) & set(input_info.keys())
        
        # --- UNUSED Redundant Check methods REMOVED to clean up codebase ---
        template_levels = {name: level for name, level in template_sections}
        for section_name in common_sections:
            # Re-find level from input_sections for accuracy
            input_level = next(level for name, level, lid in input_sections if name == section_name)
            if template_levels[section_name] != input_level:
                findings.append(Finding(category="Structure", rule_id="HEADING_LEVEL_MISMATCH", issue=f"Section '{section_name}' level mismatch", location_id=input_info[section_name][1], suggestion=f"Change to Heading {template_levels[section_name]}"))
        return findings

    def _check_table_structure(self, input_model: DocumentModel, template_model: DocumentModel) -> Tuple[int, int, List[Finding]]:
        findings = []
        table_matches = 0
        table_total_expected = 0
        
        template_sections_map = {s.heading_text.lower().strip(): s for s in template_model.sections if s.heading_level > 0}
        input_sections_map = {s.heading_text.lower().strip(): s for s in input_model.sections if s.heading_level > 0}
        
        for section_name, template_section in template_sections_map.items():
            template_table_ids = [cid for cid in template_section.content_ids if cid in template_model.tables]
            count_in_section_expected = len(template_table_ids)
            table_total_expected += count_in_section_expected
            
            if section_name in input_sections_map:
                input_section = input_sections_map[section_name]
                input_table_ids = [cid for cid in input_section.content_ids if cid in input_model.tables]
                
                # Check 1: Count Mismatch
                if len(template_table_ids) != len(input_table_ids):
                    loc_id = input_section.content_ids[0] if input_section.content_ids else "document"
                    findings.append(Finding(category="Structure", rule_id="TABLE_COUNT_MISMATCH", issue=f"Section '{section_name}' table count mismatch (expected {len(template_table_ids)}, found {len(input_table_ids)})", location_id=loc_id, location_type="section", suggestion=f"Adjust to {len(template_table_ids)} tables"))
                
                # Check 2: Structure match (up to min length)
                # We count valid matches even if total count is wrong (e.g. 2 of 3 match)
                for i in range(min(len(template_table_ids), len(input_table_ids))):
                    temp_table = template_model.tables[template_table_ids[i]]
                    inp_table = input_model.tables[input_table_ids[i]]
                    
                    if temp_table.col_count == inp_table.col_count:
                        table_matches += 1
                    else:
                        findings.append(Finding(category="Structure", rule_id="TABLE_COLUMN_MISMATCH", issue=f"Table column count mismatch in '{section_name}' table {i+1} (expected {temp_table.col_count}, found {inp_table.col_count})", location_id=input_table_ids[i], location_type="table", suggestion=f"Adjust to {temp_table.col_count} columns"))
                        
        return table_matches, table_total_expected, findings

    def _check_references(self, model: DocumentModel) -> List[Finding]:
        for section in model.sections:
            if section.heading_text and 'reference' in section.heading_text.lower():
                return []
        return [Finding(category="Content", rule_id="MISSING_REFERENCES", issue="Missing 'References' section.", location_id="document", suggestion="Please add a References section.")]

    def _check_fonts(self, model: DocumentModel) -> List[Finding]:
        """Aligning with validate_fonts implementation"""
        return validate_fonts(model)

    def _check_acronyms(self, model: DocumentModel) -> List[Finding]:
        findings = []
        acronym_section = None
        for section in model.sections:
            if section.heading_text and 'acronym' in section.heading_text.lower():
                acronym_section = section
                break
        
        if not acronym_section:
            return [Finding(category="Content", rule_id="MISSING_ACRONYMS_SECTION", issue="Missing Acronyms section", location_id="document", suggestion="Add Acronyms section")]
        
        acronyms_found = []
        for content_id in acronym_section.content_ids:
            if content_id in model.paragraphs:
                para = model.paragraphs[content_id]
                matches = re.findall(r'^([A-Z]{2,})\s*[-–—:]', para.text.strip())
                if matches: acronyms_found.append(matches[0])
        
        if len(acronyms_found) > 1 and acronyms_found != sorted(acronyms_found):
            findings.append(Finding(category="Content", rule_id="ACRONYMS_NOT_ALPHABETICAL", issue="Acronyms not alphabetical", location_id=acronym_section.id, suggestion="Sort alphabetically"))
        return findings

    def _check_sections(self, model: DocumentModel) -> List[Finding]:
        """
        Check for empty or too-brief sections.
        Flags sections that have headers but no meaningful content.
        """
        findings = []
        
        logger.info(f"DEBUG: Checking {len(model.sections)} sections for empty/brief content...")
        
        for section in model.sections:
            # Skip root section (level 0) as it's typically the document body before first heading
            if section.heading_level == 0:
                logger.debug(f"DEBUG: Skipping root section")
                continue
            
            logger.info(f"DEBUG: Analyzing section '{section.heading_text}' (Level {section.heading_level}, ID: {section.id})")
            
            # CRITICAL FIX: Determine the location_id for comment attachment
            # Use the first paragraph in the section (usually the heading itself)
            comment_location_id = section.content_ids[0] if section.content_ids else "document"
                
            # Check if section is completely empty
            if not section.content_ids:
                logger.warning(f"DEBUG: Section '{section.heading_text}' is COMPLETELY EMPTY (no content_ids)")
                findings.append(Finding(
                    category="Structure", 
                    rule_id="EMPTY_SECTION", 
                    issue=f"Section '{section.heading_text}' appears empty or too brief.", 
                    location_id=comment_location_id,  # Use first para ID for comment attachment
                    location_type="section", 
                    suggestion=f"Provide detailed content for this section in accordance with GDP standards. Section '{section.heading_text}' requires substantive explanation.",
                    priority="MEDIUM"
                ))
                continue
            
            # Check if section has content but it's too brief
            # Calculate actual text content length (excluding heading)
            total_text_length = 0
            para_count = 0
            table_count = 0
            
            for cid in section.content_ids:
                if cid in model.paragraphs:
                    para_text = model.paragraphs[cid].text.strip()
                    total_text_length += len(para_text)
                    para_count += 1
                    logger.debug(f"DEBUG:   Para {cid}: {len(para_text)} chars")
                elif cid in model.tables:
                    # Tables count as substantial content
                    total_text_length += 200  # Assume tables are meaningful
                    table_count += 1
                    logger.debug(f"DEBUG:   Table {cid}: counted as 200 chars")
            
            logger.info(f"DEBUG: Section '{section.heading_text}' total: {total_text_length} chars ({para_count} paras, {table_count} tables)")
            
            # Flag sections with less than 100 characters of content (more aggressive threshold)
            MIN_SECTION_LENGTH = 100
            if total_text_length < MIN_SECTION_LENGTH and total_text_length > 0:
                logger.warning(f"DEBUG: Section '{section.heading_text}' is TOO BRIEF ({total_text_length} < {MIN_SECTION_LENGTH})")
                findings.append(Finding(
                    category="Structure", 
                    rule_id="BRIEF_SECTION", 
                    issue=f"Section '{section.heading_text}' appears too brief (only {total_text_length} characters).", 
                    location_id=comment_location_id,  # Use first para ID for comment attachment
                    location_type="section", 
                    suggestion=f"Expand this section with detailed content in accordance with GDP standards. Section '{section.heading_text}' requires more comprehensive explanation.",
                    priority="MEDIUM"
                ))
        
        logger.info(f"DEBUG: Found {len(findings)} empty/brief sections")
        return findings



    def _check_tables(self, model: DocumentModel) -> List[Finding]:
        findings = []
        for table_id, table in model.tables.items():
            if table.content and all(not any(cell.strip() for cell in row) for row in table.content):
                findings.append(Finding(category="Structure", rule_id="BLANK_TABLE", issue="Table is completely empty.", location_id=table_id, location_type="table", suggestion="Please remove the table or fill it with relevant content."))
                continue

            if table.has_empty_cells:
                findings.append(Finding(category="Structure", rule_id="EMPTY_CELLS", issue="Table contains empty cells.", location_id=table_id, location_type="table", suggestion="Please populate empty cells with relevant data or a justification phrase."))
            
            # Enhanced multi-page table detection
            is_header_repeated = table.metadata.get("is_header_repeated", False)
            likely_multipage = False
            
            # Check if table likely spans multiple pages
            if table.row_count > 15:
                likely_multipage = True
            elif table.row_count > 10:
                # Estimate content size
                total_content_length = sum(len(str(cell)) for row in table.content for cell in row)
                # If average cell content is long, likely spans pages
                avg_cell_length = total_content_length / (table.row_count * table.col_count) if table.row_count * table.col_count > 0 else 0
                if avg_cell_length > 50 or total_content_length > 1500:
                    likely_multipage = True
            
            if likely_multipage and not is_header_repeated:
                findings.append(Finding(
                    category="Formatting",
                    rule_id="TABLE_HEADER_NOT_REPEATED",
                    issue=f"Table with {table.row_count} rows likely spans multiple pages without repeated headers",
                    location_id=table_id,
                    location_type="table",
                    suggestion="Enable 'Repeat Header Rows' in table properties"
                ))
            
            if table.headers:
                weak = [h for h in table.headers if len(h) <= 1 or "column" in h.lower()]
                if weak: findings.append(Finding(category="Structure", rule_id="WEAK_TABLE_HEADER", issue=f"Weak headers found: {', '.join(weak)}", location_id=table_id, location_type="table", suggestion="Use descriptive headers"))
                
                if lower and len(lower) == len([h for h in table.headers if h]):
                    findings.append(Finding(category="Style", rule_id="LOWERCASE_TABLE_HEADER", issue="Lowercase headers", location_id=table_id, location_type="table", suggestion="Use Title Case"))
        return findings

    def _check_dates_in_document(self, model: DocumentModel) -> List[Finding]:
        findings = []
        for para_id, para in model.paragraphs.items():
            # Strict DD/MM/YYYY regex
            # DD: 01-31, MM: 01-12, YYYY: 4 digits
            date_pattern = r"\b(0[1-9]|[12]\d|3[01])/(0[1-9]|1[0-2])/(\d{4})\b"
            
            # Find all potential dates using a broad pattern to catch incorrectly formatted ones
            # Matches anything looking like d/m/y or d-m-y
            matches = re.finditer(r"\b(?:\d{1,2}[/-]\d{1,2}[/-]\d{2,4})\b", para.text)
            
            for match in matches:
                date_str = match.group()
                
                # Check against strict DD/MM/YYYY
                if not re.match(date_pattern, date_str):
                     findings.append(Finding(category="Content", rule_id="INVALID_DATE_FORMAT", issue=f"Date '{date_str}' is not in the required DD/MM/YYYY format.", location_id=para_id, suggestion="Please use the strict DD/MM/YYYY format."))
                else:
                    # Further validate logic (e.g. 30/02/2023 would pass regex but fail calendar)
                    try:
                        # Parse as DD/MM/YYYY
                        d = datetime.strptime(date_str, "%d/%m/%Y")
                        if d.year < 1900 or d.year > 2100:
                             findings.append(Finding(category="Content", rule_id="INVALID_DATE_VALUE", issue=f"Date '{date_str}' year out of reasonable range.", location_id=para_id, suggestion="Please verify the year."))
                    except ValueError:
                         findings.append(Finding(category="Content", rule_id="INVALID_DATE_VALUE", issue=f"Date '{date_str}' is invalid (e.g. Feb 30).", location_id=para_id, suggestion="Please correct the date."))
        return findings

    def _check_whitespace_in_document(self, model: DocumentModel) -> List[Finding]:
        findings = []
        for para_id, para in model.paragraphs.items():
            if "  " in para.text:
                findings.append(Finding(category="Formatting", rule_id="MULTIPLE_SPACES", issue="Multiple spaces found", location_id=para_id, suggestion="Use single space"))
            for line in para.text.split('\n'):
                if line != line.rstrip():
                    findings.append(Finding(category="Formatting", rule_id="TRAILING_SPACES", issue="Trailing spaces found", location_id=para_id, suggestion="Remove trailing spaces"))
                    break
        return findings

    def insert_comments(self, file_path: str, findings: List[Finding], output_path: str, template_name: str = None):
        """
        Insert comments into the Word document at exact error locations.
        Enables color highlighting by default to provide visual cues for consolidated errors.
        """
        import docx
        from docx.shared import Pt
        
        # Load the document (using the bookmarked version if available)
        doc = docx.Document(file_path)
        enable_highlight = True # Always enable highlighting for clarity
        logger.info(f"VISUAL HIGHLIGHTING ENABLED for template: {template_name if template_name else 'Unknown'}")
        
        # DEBUG: Log findings to verify priority presence
        logger.info(f"Processing {len(findings)} findings for comments...")
        for i, f in enumerate(findings[:5]): # Log first 5
             issue = f.issue if hasattr(f, 'issue') else f.get('issue', '')
             prio = getattr(f, 'priority', 'N/A') if hasattr(f, 'priority') else f.get('priority', 'N/A')
             logger.info(f"Finding {i}: Rule={getattr(f, 'rule_id', 'N/A')}, Priority={prio}, IssueStart='{issue[:30]}...'")

        # ---------------------------------------------------------
        # CLEANUP: Remove any existing highlighting (Code Omitted for Brevity)
        # ---------------------------------------------------------
        from docx.oxml.ns import qn
        def remove_run_highlight(run):
            rPr = run.element.rPr
            if rPr and rPr.find(qn('w:highlight')) is not None:
                rPr.remove(rPr.find(qn('w:highlight')))
        for p in doc.paragraphs: [remove_run_highlight(r) for r in p.runs]
        for t in doc.tables: 
            for r in t.rows: 
                for c in r.cells: 
                    for p in c.paragraphs: [remove_run_highlight(r) for r in p.runs]
        # ---------------------------------------------------------

        # Build sorting map (Code Omitted for Brevity - unchanged)
        sort_map = {}
        curr_idx = 0
        p_count = 0
        t_count = 0
        for element in doc.element.body:
             if isinstance(element, CT_P):
                 sort_map[f"para_{p_count}"] = curr_idx
                 p_count += 1
             elif isinstance(element, CT_Tbl):
                 sort_map[f"table_{t_count}"] = curr_idx
                 t_count += 1
             curr_idx += 1

        def get_global_sort_id(f):
             loc = f.location_id
             if loc in sort_map: return sort_map[loc]
             if f.location_type == 'table' or loc.startswith('table_'):
                 base = loc.split('_r')[0]
                 if base in sort_map: return sort_map[base]
             if f.location_type in ['document', 'structure', 'section']: return -1
             return 99999

        findings.sort(key=lambda f: (get_global_sort_id(f), f.location_id))
        
        findings_map = defaultdict(list)
        for f in findings: findings_map[f.location_id].append(f)
        
        # ... (Rest of iteration logic remains checking sort_map/findings_map) ...
        # (For simplicity of this edit, I am returning to the main iteration logic but asking to keep it unchanged 
        #  via context, but I need to be careful not to delete the iteration code. 
        #  I will only replace the top part of the function and the _attach_comment_with_priority helper).
        
        # RE-IMPLEMENTING ITERATION LOGIC TO BE SAFE BECAUSE "ReplacementContent" replaces the whole block
        def iter_block_elements(parent):
            for child in parent.iterchildren():
                if isinstance(child, CT_P): yield ('paragraph', docx.text.paragraph.Paragraph(child, doc))
                elif isinstance(child, CT_Tbl): yield ('table', docx.table.Table(child, doc))
                elif child.tag.endswith('sdt'):
                    content = child.find(qn('w:sdtContent'))
                    if content: yield from iter_block_elements(content)

        body_elements = list(iter_block_elements(doc.element.body))
        para_count = 0
        table_count = 0
        last_paragraph = None
        
        for i, (elem_type, elem) in enumerate(body_elements):
            if elem_type == 'paragraph':
                para_id = f"para_{para_count}"
                last_paragraph = elem
                if para_id in findings_map:
                    for f in findings_map[para_id]:
                        self._attach_comment_with_priority(elem, f, enable_highlight=enable_highlight)
                para_count += 1
            elif elem_type == 'table':
                table_id = f"table_{table_count}"
                table_findings = []
                if table_id in findings_map: table_findings.extend(findings_map[table_id])
                
                cell_findings = {}
                prefix = table_id + "_"
                for key in findings_map:
                    if key.startswith(prefix) and "_r" in key and "_c" in key:
                         cell_findings[key] = findings_map[key]
                
                if table_findings:
                    cell_rule_ids = {f.rule_id if hasattr(f, 'rule_id') else f.get('rule_id') for fs in cell_findings.values() for f in fs}
                    target = last_paragraph
                    if not target:
                        for j in range(i+1, len(body_elements)):
                            if body_elements[j][0] == 'paragraph': target = body_elements[j][1]; break
                    if target:
                        for f in table_findings:
                            rid = f.rule_id if hasattr(f, 'rule_id') else f.get('rule_id')
                            if rid in cell_rule_ids: continue
                            self._attach_comment_with_priority(target, f, enable_highlight=enable_highlight)

                for cell_loc_id, c_findings in cell_findings.items():
                    try:
                        parts = cell_loc_id.split('_')
                        r_idx = int([p[1:] for p in parts if p.startswith('r')][0])
                        c_idx = int([p[1:] for p in parts if p.startswith('c')][0])
                        cell = elem.cell(r_idx, c_idx)
                        if cell.paragraphs:
                            for f in c_findings:
                                self._attach_comment_with_priority(cell.paragraphs[0], f, prefix=f"[Row {r_idx+1}, Col {c_idx+1}]", enable_highlight=enable_highlight)
                    except: pass
                table_count += 1
        
        if body_elements and "document" in findings_map:
             # Anchor document-wide findings (like MISSING_TOC) to the very first paragraph
             target_para = None
             for etype, elem in body_elements:
                 if etype == 'paragraph':
                     target_para = elem
                     break
             
             if target_para:
                 for f in findings_map["document"]:
                     self._attach_comment_with_priority(target_para, f, prefix="[Document Level] ", enable_highlight=enable_highlight)

        doc.core_properties.author = "GDP Doc Reviewer"
        doc.core_properties.last_modified_by = "GDP Doc Reviewer"
        doc.save(output_path)
        self.generate_validation_report(file_path, findings, output_path)

    def _attach_comment_with_priority(self, element, finding, prefix="", enable_highlight=False):
        """Helper to attach a comment with priority labeling and color highlighting."""
        from docx.enum.text import WD_COLOR_INDEX
        
        # Priority mapping to Word Colors
        PRIORITY_WORD_COLOR = {
            "HIGH": WD_COLOR_INDEX.RED,
            "MEDIUM": WD_COLOR_INDEX.YELLOW,
            "LOW": WD_COLOR_INDEX.BRIGHT_GREEN
        }
        # Weighting for color overrides
        COLOR_WEIGHTS = {WD_COLOR_INDEX.RED: 3, WD_COLOR_INDEX.YELLOW: 2, WD_COLOR_INDEX.BRIGHT_GREEN: 1}

        try:
            # 1. Safely extract priority and text
            if hasattr(finding, 'issue'):
                 issue_text = str(finding.issue).strip()
                 sugg_text = str(finding.suggestion).strip()
                 priority = str(getattr(finding, 'priority', 'MEDIUM')).upper()
            else: # Dictionary case
                 issue_text = str(finding.get('issue', '')).strip()
                 sugg_text = str(finding.get('suggestion', '')).strip()
                 priority = str(finding.get('priority', 'MEDIUM')).upper()
            
            # 2. Handle Location Prefix Injection (e.g. [Row 1, Col 2])
            # Desired outcome: [Priority is X] [Row 1, Col 2] Observation: ...
            
            import re
            
            print(f"DEBUG_ATTACH_INPUT: Issue='{issue_text[:30]}...' Sugg='{sugg_text[:30]}...'")

            # CLEAN AND REBUILD (DETERMINISTIC)
            # We always clean the raw text first, then wrap it in the latest standard format.
            # This is the single source of truth for the Word document comment body.
            
            clean_issue = clean_finding_text(issue_text)
            clean_sugg = clean_finding_text(sugg_text)
            
            # 2a. FAILSAFE: PANIC MODE
            # If the cleaner failed (e.g. infinite loop of junk), force strip known strings
            # 2a. FAILSAFE: PANIC MODE (Case Insensitive)
            # If the cleaner failed, force strip known strings using Regex
            # Check for [Priority...] Observation: pattern anywhere
            if re.search(r'\[priority', clean_issue, re.IGNORECASE) and re.search(r'observation\s*:', clean_issue, re.IGNORECASE):
                 # Split on the LAST occurrence of "Observation:" to catch nested dupes
                 parts = re.split(r'observation\s*:', clean_issue, flags=re.IGNORECASE)
                 if len(parts) > 1:
                     clean_issue = parts[-1].strip()

            if re.search(r'recommendation\s*:', clean_sugg, re.IGNORECASE):
                 parts = re.split(r'recommendation\s*:', clean_sugg, flags=re.IGNORECASE)
                 if len(parts) > 1:
                     clean_sugg = parts[-1].strip()
            
            # Use title case for label
            display_priority = priority.upper().title() if priority else "Medium"
            p_label = f"[Priority is {display_priority}]"
            
            # Construct final body with double newline
            # Result: [Priority is Low] Observation: Heading ends with a period.
            if prefix: # Location prefix (e.g. [Row 1, Col 1])
                comment_body = f"{p_label} Observation: {prefix} {clean_issue} \n\n Recommendation: {clean_sugg}".strip()
            else:
                comment_body = f"{p_label} Observation: {clean_issue} \n\n Recommendation: {clean_sugg}".strip()

            # 3. FINAL REGEX DEDUPLICATION check
            # Pattern matches: [Priority...] Observation: <Space/Junk> [Priority...] Observation:
            # We explicitly match the pattern TWICE instead of using backreference \1 to allow for Case/Content diffs
            # e.g. [Priority is Low] vs [Priority is LOW]
            
            dedupe_pattern = r'^(\[[^\]]+\]\s*Observation:)\s*(\[[^\]]+\]\s*Observation:)'
            match = re.match(dedupe_pattern, comment_body, flags=re.IGNORECASE)
            if match:
                 # We found a double header. Keep only the first one (group 1)
                 # and the rest of the string after the second header
                 first_header = match.group(1)
                 full_match_len = match.end()
                 rest_of_body = comment_body[full_match_len:].strip()
                 comment_body = f"{first_header} {rest_of_body}"
            
            print(f"DEBUG_ATTACH_BODY: '{comment_body[:60]}...'")

            # 3. Add to Word Document
            add_comment_to_paragraph(element, comment_body)
            
            # 4. Color Highlighting
            if enable_highlight:
                word_color = PRIORITY_WORD_COLOR.get(priority)
                if word_color:
                    target_paras = [element] if hasattr(element, 'runs') else getattr(element, 'paragraphs', [])
                    for p in target_paras:
                        for run in p.runs:
                            # Apply highlight if none exists or if new color is higher weight (more critical)
                            ch = run.font.highlight_color
                            if ch is None or COLOR_WEIGHTS.get(word_color, 0) > COLOR_WEIGHTS.get(ch, 0):
                                run.font.highlight_color = word_color
                                
        except Exception as e:
            logger.error(f"Error in _attach_comment_with_priority: {e}", exc_info=True)
            # Minimal fallback
            try:
                msg = f"[System Alert] {getattr(finding, 'issue', str(finding))}"
                add_comment_to_paragraph(element, msg)
            except: pass

    def generate_validation_report(self, input_path: str, findings: List[Finding], output_docx_path: str):
        """Generates a JSON summary report of the validation."""
        report_path = output_docx_path.replace(".docx", "_report.json")
        
        summary = {
            "document_name": os.path.basename(input_path),
            "review_date": datetime.now().isoformat(),
            "total_findings": len(findings),
            "findings_by_category": {},
            "findings_by_severity": {"BLOCKING": 0, "HIGH": 0, "MEDIUM": 0, "LOW": 0},
            "critical_violations": []
        }
        
        for f in findings:
            cat = f.category
            summary["findings_by_category"][cat] = summary["findings_by_category"].get(cat, 0) + 1
            
            sev = f.priority if f.priority in summary["findings_by_severity"] else "MEDIUM"
            summary["findings_by_severity"][sev] += 1
            
            if sev in ["BLOCKING", "HIGH"]:
                summary["critical_violations"].append({
                    "rule_id": f.rule_id,
                    "issue": f.issue,
                    "location": f.location_id
                })
                
        try:
            with open(report_path, 'w') as f:
                json.dump(summary, f, indent=2)
        except Exception as e:
            logging.error(f"Failed to write validation report: {e}")


# ============================================================================
# VALIDATOR
# ============================================================================

DEFAULT_GUIDELINES = [
    {"id": "GRAMMAR_PHRASING", "text": "Ensure grammar is correct and phrasing is clear.", "keywords": ["grammar", "clarity"]},
    {"id": "MISLEADING_CONTENT", "text": "Check for misleading, contradictory, or irrelevant content.", "keywords": ["content", "relevance"]},
    {"id": "SEMANTIC_EMPTY", "text": "Check if section content is semantically empty (e.g. no meaningful info). Do NOT flag 'TBD' or 'Pending' as this is handled by strict checks.", "keywords": ["empty"]}
]

@dataclass
class DocumentState:
    defined_acronyms: Set[str] = field(default_factory=set)
    seen_acronyms: Set[str] = field(default_factory=set)
    document_authors: Set[str] = field(default_factory=set)
    document_reviewers: Set[str] = field(default_factory=set)
    document_approvers: Set[str] = field(default_factory=set)

DEFAULT_GEMINI_MODEL = 'gemini-2.5-flash'

class LLMFinding(BaseModel):
    category: Annotated[str, Field(pattern=r'^(Content|Structure|Style|Formatting|Authorship|Spelling|Grammar)$')]
    rule_id: str
    issue: Annotated[str, Field(min_length=1, max_length=400)]
    location_id: Annotated[str, Field(pattern=r'^(para_\d+|table_\d+.*|document|P\d+|unknown)$')]
    suggestion: Annotated[str, Field(min_length=1, max_length=200)]
    confidence: Optional[Annotated[float, Field(ge=0.0, le=1.0)]] = None
    evidence: Optional[Annotated[str, Field(max_length=500)]] = None

class LLMClient:
    """Wrapper for Azure OpenAI calls with robust JSON extraction and schema validation."""
    def __init__(self, api_key: str, azure_endpoint: str, api_version: str, deployment_name: str):
        self.client = AzureOpenAI(
            api_key=api_key,
            api_version=api_version,
            azure_endpoint=azure_endpoint
        )
        self.deployment_name = deployment_name

    def _clean_model_output(self, text: str) -> str:
        if not text:
            return ""
        s = text.strip()
        if s.startswith("```"):
            parts = s.split("```", 2)
            if len(parts) >= 3:
                s = parts[1].strip()
            else:
                s = s.strip('`').strip()
        s = re.sub(r'^\s*(json|output|answer)\s*[:\-]*\s*', '', s, flags=re.I)
        first = min([i for i in (s.find('['), s.find('{')) if i != -1], default=None)
        if first is not None:
            s = s[first:]
        last = max(s.rfind(']'), s.rfind('}'))
        if last != -1:
            s = s[:last+1]
        return s.strip()

    def _validate_and_parse(self, raw_text: str) -> List[Dict[str, Any]]:
        """Try to robustly parse raw_text into a list of LLMFinding-like dicts."""
        cleaned = self._clean_model_output(raw_text)
        if not cleaned:
            return []

        try:
            parsed = json.loads(cleaned)
        except JSONDecodeError:
            last = max(cleaned.rfind(']'), cleaned.rfind('}'))
            if last != -1:
                try:
                    parsed = json.loads(cleaned[:last+1])
                except Exception:
                    return []
            else:
                return []
        if isinstance(parsed, dict):
            parsed = [parsed]
        if not isinstance(parsed, list):
            return []

        validated = []
        for item in parsed:
            try:
                obj = LLMFinding(**item)
                validated.append(obj.dict())
            except Exception as ve:
                if isinstance(item.get("location_id"), str):
                    item["location_id"] = re.sub(r'\s+', '', item["location_id"])
                try:
                    obj = LLMFinding(**item)
                    validated.append(obj.dict())
                except Exception:
                    logger.debug("LLM item failed schema validation: %s", item)
                    continue
        return validated

    def generate_json(self, prompt: str, schema_description: str = "", retries: int = 3, retry_on_invalid: bool = True, system_prompt: str = None) -> List[Dict[str, Any]]:
        """Call the model and return a list of validated dicts matching LLMFinding."""
        full_prompt = f"{prompt}\n\nIMPORTANT: Return ONLY a valid JSON array following the schema. No markdown."
        
        messages = []
        if system_prompt:
             messages.append({"role": "system", "content": system_prompt})
        messages.append({"role": "user", "content": full_prompt})

        for attempt in range(max(1, retries)):
            try:
                response = self.client.chat.completions.create(
                    model=self.deployment_name,
                    messages=messages,
                    temperature=0,
                    max_tokens=12000,
                    top_p=1.0,
                    frequency_penalty=0.0,
                    presence_penalty=0.0
                )
                
                raw = response.choices[0].message.content
                results = self._validate_and_parse(raw)
                if results:
                    return results
                
                if retry_on_invalid and attempt < retries - 1:
                    logger.debug("LLM output invalid JSON/schema; attempting reformat (attempt %d). Raw head: %s", attempt, (raw or "")[:400])
                    reformat_prompt = f"{full_prompt}\n\nThe previous output was not valid JSON. Please reformat your previous answer and return EXACTLY a valid JSON array matching the schema."
                    
                    reform_messages = []
                    if system_prompt:
                         reform_messages.append({"role": "system", "content": system_prompt})
                    reform_messages.append({"role": "user", "content": reformat_prompt})
                    
                    response2 = self.client.chat.completions.create(
                        model=self.deployment_name,
                        messages=reform_messages,
                        temperature=0,
                        max_tokens=12000
                    )
                    raw2 = response2.choices[0].message.content
                    results = self._validate_and_parse(raw2)
                    if results:
                        return results
                
                # Exponential backoff
                sleep_time = 0.5 * (2 ** attempt)
                time.sleep(sleep_time)
                
            except Exception as e:
                logger.warning("LLM call failed attempt %d: %s", attempt, e)
                sleep_time = 1.0 * (2 ** attempt) # Slightly longer sleep on exception
                time.sleep(sleep_time)
        return []

def llm_safe_text(text: str) -> str:
    cleaned = []
    
    for line in text.splitlines():
        line_clean = line.strip()
        if not line_clean:
            continue
            
        # REMOVE URLs
        if re.search(r'https?://|www\.', line_clean, re.I):
            continue
            
        # REMOVE tables / revision rows / strict metadata lines
        # Matches "Version 1.0", "Date: ...", "Author ...", "Reviewer..."
        if re.search(r'^\s*(version|date|author|reviewer|approver)\b.*[\d\w]', line_clean, re.I):
            continue
            
        # REMOVE purely numeric/symbolic lines
        if re.match(r'^[\d\s\.\-,;:\(\)\[\]]+$', line_clean):
            continue
            
        # REMOVE lines that are purely list items without content
        if re.match(r'^[\d\w][\)\.]\s*$', line_clean):
            continue

        # PRESERVE lines with acronyms if they look like sentences
        # Only remove if line is VERY short and looks like garbage
        if len(line_clean.split()) < 2:
             # Remove single tokens that are short OR look like codes (contain digits)
             # "Introduction" (len 12) -> Keep
             # "A" (len 1) -> Remove
             # "X12" (len 3, digits) -> Remove
             # "SOP" (len 3, no digits) -> Keep
             if len(line_clean) < 3 or (len(line_clean) < 6 and any(c.isdigit() for c in line_clean)):
                 continue

        cleaned.append(line)
        
    return "\n".join(cleaned)

    
def dedupe_findings(findings: List[Finding]) -> List[Finding]:
    """
    Consolidates multiple findings.
    Includes AGGRESSIVE filtering to remove duplicate/redundant rules.
    """
    from collections import defaultdict
    import re
    
    seen_hashes = set()
    cleaned_list = []
    
    priority_order = {"BLOCKING": 0, "HIGH": 1, "MEDIUM": 2, "LOW": 3}

    for f in findings:
        # --- AGGRESSIVE FILTERS ---
        # 1. Kill "DATE_FORMAT_ERROR" (Medium/LLM) in favor of "INVALID_DATE_FORMAT" (High/Regex)
        if f.rule_id == "DATE_FORMAT_ERROR":
            continue

        # 2. Kill "MISSING_CONTEXT" if it talks about TBD (duplicates TBD_PLACEHOLDER)
        if f.rule_id == "MISSING_CONTEXT" and re.search(r'\bTBD\b', f.issue, re.IGNORECASE):
            continue

        # 3. Kill "TBD_PLACEHOLDER" if priority is NOT HIGH (legacy check leftovers)
        if f.rule_id == "TBD_PLACEHOLDER" and f.priority != "HIGH":
             continue

        # 4. Kill "GRAMMAR_SPELLING" (Medium/LLM) in favor of "MISSPELLED_WORD" (Low/Static)
        # User explicitly requested removal of repeated spell check findings
        if f.rule_id == "GRAMMAR_SPELLING":
             continue

        # Normal cleaning
        f.issue = clean_finding_text(f.issue)
        f.suggestion = clean_finding_text(f.suggestion)
        f.suggestion = enhance_suggestion(f) 

        # Unique signature: (Location, Rule, Issue Text)
        sig = (f.location_id, f.rule_id, f.issue)
        if sig in seen_hashes:
            continue
        seen_hashes.add(sig)
        cleaned_list.append(f)

    # Sort by priority then location for clean output
    cleaned_list.sort(key=lambda x: (priority_order.get(x.priority.upper(), 2), x.location_id))
            
    return cleaned_list
            
    return cleaned_list


def enhance_suggestion(finding: Finding) -> str:
    """Enhance suggestions with specific, actionable recommendations from templates."""
    rule_id = finding.rule_id
    
    # Load templates from data directory
    templates = {}
    try:
        data_dir = os.path.join(os.path.dirname(__file__), 'data')
        tpl_path = os.path.join(data_dir, 'comment_templates.json')
        if os.path.exists(tpl_path):
            with open(tpl_path, 'r', encoding='utf-8') as f:
                templates = json.load(f)
    except Exception:
        pass

    # Use template if available, otherwise fallback to hardcoded or existing
    if rule_id in templates:
        suggestion = templates[rule_id]
    else:
        # Check if the existing suggestion is already specific/detailed
        # If it contains "pt" or "standard" or "Reference", we keep it
        curr_sugg = finding.suggestion or ""
        if any(keyword in curr_sugg.lower() for keyword in ["pt", "title case", "period", "standard"]) and rule_id in ["INCORRECT_FONT_SIZE", "HEADING_CASE_ISSUE", "HEADING_WITH_PERIOD", "MISSING_PERIOD_BULLET"]:
             suggestion = curr_sugg
        else:
            fallbacks = {
                "INCORRECT_FONT_SIZE": "Select the text and change font size to the required GDP standard (H1: 24pt, H2: 18pt, H3: 16pt, Normal: 12pt).",
                "INCORRECT_FONT": "Select the text and change font to Arial (Home → Font → Arial)",
                "MISSING_PERIOD": "Add a period (.) at the end of the bullet point",
                "EMPTY_SECTION": "Add relevant content to the section or provide a clear justification explaining its omission (e.g., 'This section is not applicable for this project scope')",
                "PERSONAL_REFERENCE": "Rewrite using third person (e.g., 'your system' → 'the system') or passive voice (e.g., 'you must' → 'it is required to')",
            }
            suggestion = fallbacks.get(rule_id, curr_sugg)
    # Remove prefix if present to allow unified adding later
    suggestion = clean_finding_text(suggestion or '')
    return suggestion


def consolidate_group(group: List[Finding], rule_id: str, base_loc: str, loc_type: str) -> Finding:
    """Consolidate multiple similar findings into one with detailed information."""
    count = len(group)
    first = group[0]
    
    # Extract specific details from all findings
    details = []
    for f in group:
        if '_r' in f.location_id and '_c' in f.location_id:
            # Extract row and column
            parts = f.location_id.split('_')
            for i, part in enumerate(parts):
                if part.startswith('r') and i+1 < len(parts) and parts[i+1].startswith('c'):
                    row = part[1:]
                    col = parts[i+1][1:]
                    details.append(f"R{int(row)+1}C{int(col)+1}")
                    break
    
    # Create consolidated issue and suggestion
    if rule_id == "EMPTY_CELLS":
        cells_list = ", ".join(details[:5])  # Show first 5
        more = f" and {len(details)-5} more" if len(details) > 5 else ""
        issue = f"Table contains {count} empty cell{'s' if count > 1 else ''} at: {cells_list}{more}"
        suggestion = f"Populate all {count} empty cells with relevant data or a specific justification phrase. Example: 'Not applicable for this project segment'"
        
    elif rule_id == "INVALID_DATE_FORMAT":
        # Extract dates from issues
        dates = []
        for f in group:
            match = re.search(r"'([^']+)'", f.issue)
            if match:
                dates.append(match.group(1))
        dates_list = ", ".join(dates[:5])
        more = f" and {len(dates)-5} more" if len(dates) > 5 else ""
        issue = f"Found {count} date{'s' if count > 1 else ''} with incorrect format: {dates_list}{more}"
        suggestion = f"Convert all {count} dates to DD/MM/YYYY format with 4-digit year. Example: '05/21/25' → '21/05/2025' (if May 21, 2025)"
        
    elif rule_id == "INCORRECT_FONT_SIZE":
        issue = f"Found {count} paragraph{'s' if count > 1 else ''} with incorrect font size"
        suggestion = f"Select all {count} paragraphs and change font size to 12pt (Home → Font Size → 12) (or correct Heading size)."
        
    elif rule_id == "PERSONAL_REFERENCE" or rule_id == "PERSONAL_REFERENCE_TABLE":
        issue = f"Found {count} instance{'s' if count > 1 else ''} of personal pronouns in {base_loc}"
        suggestion = f"Rewrite all {count} instances using third person or passive voice. Example: 'your system' → 'the system', 'you must' → 'it is required to'"
        
    else:
        # Generic consolidation
        issue = f"Found {count} {rule_id.replace('_', ' ').lower()} issue{'s' if count > 1 else ''} in {base_loc}"
        suggestion = enhance_suggestion(first)
    
    # Create consolidated finding
    return Finding(
        category=first.category,
        rule_id=rule_id,
        issue=issue,
        location_type=loc_type,
        location_id=base_loc,
        suggestion=suggestion,
        priority=first.priority,
        confidence=first.confidence
    )

class Validator:
    def __init__(self, config=None):
        # Merge provided config with DEFAULT_CONFIG
        self.config = DEFAULT_CONFIG.copy()
        if config:
            self.config.update(config)
            
        self.prompt_builder = PromptBuilder()
        try:
            self.llm_client = LLMClient(
                api_key=self.config.get("AZURE_OPENAI_API_KEY") or os.environ.get("AZURE_OPENAI_KEY"),
                azure_endpoint=self.config.get("AZURE_OPENAI_ENDPOINT"),
                api_version=self.config.get("AZURE_OPENAI_API_VERSION", "2024-02-15-preview"),
                deployment_name=self.config.get("AZURE_DEPLOYMENT_NAME") or self.config.get("AZURE_OPENAI_DEPLOYMENT", "gpt-4o")
            )
        except Exception as e:
            logger.warning(f"LLM Client init failed: {e}")
            self.llm_client = None

        # Noise Filter Set
        self.FALSE_ACRONYMS = {
            "USER", "DOCUMENT", "CONFIGURATION", "MATRIX",
            "ITEMS", "DATE", "VERSION", "FIRST", "SECOND",
            "REFERENCES", "KEY", "CONTROL", "DATA", "SYSTEM", 
            "INPUT", "OUTPUT", "REPORT", "PAGE", "SECTION"
        }

    def validate_document(self, model: DocumentModel, template_model: DocumentModel = None) -> List[Finding]:
        findings = []
        
        # 1. Deterministic Checks (SOURCE OF TRUTH)
        logger.info("Running deterministic checks...")
        # Call directly from module scope
        findings.extend(run_all_validators(model, template_model, self.config))
        
        # 2. STRICT Syntax Pipeline (Grammar, Spelling, Acronyms)
        logger.info("Running strict syntax pipeline...")
        findings.extend(self._run_strict_syntax_pipeline(model))

        # 3. LLM Checks (Semantic) - ONLY for semantic/logical rules
        semantic_guidelines = [
            {"id": "LOGICAL_CONSISTENCY", "text": "Ensure content is logically consistent across sections."},
            {"id": "MISSING_CONTEXT", "text": "Flag statements lacking context (e.g. 'Maintain ISO 27001 compliance' without specifying monitoring/audits, or 'Follow ISO annexures' without specifying which ones, or missing OS/Browser versions)."},
            {"id": "AMBIGUOUS_REQUIREMENT", "text": "Flag ambiguous terms lacking measurable criteria (e.g., 'seamless', 'intuitive', 'where applicable', 'as needed', 'if required'). Requirement must have clear metrics/Downtime limits/User impact thresholds."},
            {"id": "GDP_CLARITY", "text": "Ensure clear, distinct, and precise language per GDP."}
        ]
        
        findings.extend(self._run_llm_pipeline(model, guidelines=semantic_guidelines))
            
        # 4. Deduplicate results (Merge/Cleanup duplicates)
        logger.info(f"Deduplicating {len(findings)} findings...")
        findings = dedupe_findings(findings)

        # 5. Final Gateway: ENSURE UI CONSISTENCY (Force Labeling & Standard Prefixes)
        # This is the single source of truth for how findings appear in Web UI and Word.
        import re
        for f in findings:
            # A. Priority Resolution
            p = get_priority(f.rule_id, f.category)
            if p: f.priority = p
            if not f.priority: f.priority = "MEDIUM"
            f.priority = str(f.priority).upper()

            # B. Atomic Clean ONLY
            # We store RAW text in Finding objects. 
            # Prefixes are added only at the presentation layer (Word export / UI display).
            f.issue = clean_finding_text(f.issue)
            f.suggestion = clean_finding_text(f.suggestion)

        return findings

    def _run_strict_syntax_pipeline(self, model: DocumentModel) -> List[Finding]:
        """
        Runs the specialized strict prompt for Grammar, Spelling, and Acronyms.
        """
        findings = []
        chunks = self._create_chunks(model)
        
        # Initialize Acronym Manager to get known acronym list
        acronym_mgr = AcronymManager()
        # Combine whitelisted + common + document defined acronyms
        # (Assuming document defined acronyms could be extracted, but for now use static lists)
        known_acronyms = list(acronym_mgr.whitelist.union(acronym_mgr.common_acronyms).union(acronym_mgr.MONTH_ABBREVS).union(acronym_mgr.pharma_acronyms))
        # Add corporate acronyms keys
        known_acronyms.extend(acronym_mgr.corporate_acronyms.keys())
        known_list_str = ", ".join(sorted(list(set(known_acronyms))))
        
        # delayed import to avoid circular dependency
        from models import VALIDATION_SYSTEM_PROMPT, VALIDATION_USER_PROMPT_TEMPLATE
        
        # Inject dynamic acronyms into System Prompt
        system_prompt_dynamic = f"{VALIDATION_SYSTEM_PROMPT}\n\nALLOWED ACRONYMS (Whitelisted):\n{known_list_str[:6000]}"

        with ThreadPoolExecutor(max_workers=3) as executor:
            future_to_chunk = {}
            for chunk in chunks:
                safe_text = llm_safe_text(chunk['chunk_text'])
                if not safe_text.strip(): continue
                
                # Construct User Prompt
                user_prompt = VALIDATION_USER_PROMPT_TEMPLATE.replace("{{TEXT}}", safe_text)
                
                future_to_chunk[executor.submit(
                    self.llm_client.generate_json, 
                    prompt=user_prompt, 
                    system_prompt=system_prompt_dynamic,
                    schema_description="Strict Syntax Violations"
                )] = chunk

            for future in as_completed(future_to_chunk):
                try:
                    raw_findings = future.result()
                    if not raw_findings: continue
                    
                    for item in raw_findings:
                        # Validate item has required fields
                        if 'issue' not in item: continue
                        
                        f = Finding(
                            category=item.get('category', 'Grammar'),
                            rule_id=item.get('rule_id', 'STRICT_CHECK'),
                            issue=item.get('issue', 'Issue found'),
                            location_id=item.get('location_id', 'unknown'),
                            suggestion=item.get('suggestion', ''),
                            confidence=float(item.get('confidence', 0.9))
                        )
                        # Fix location if generic
                        if f.location_id == 'unknown':
                            # Try to fallback to section ID/Paragraph ID from chunk metadata?
                            # The strict prompt asks for "para_X". Ideally it respects that.
                            # If not, we might map to the chunk's main location.
                            chunk_meta = future_to_chunk[future]['metadata']
                            f.location_id = chunk_meta.get('location_id', 'document')
                            
                        # Skip manual prefixing, let the final gateway handle it
                        findings.append(f)
                        
                except Exception as e:
                    logger.error(f"Strict pipeline chunk failed: {e}")

        return findings

    def _run_llm_pipeline(self, model: DocumentModel, guidelines: List[Dict]) -> List[Finding]:
        findings = []
        chunks = self._create_chunks(model)
        
        # optimization: Increase workers to maximize throughput against Azure OpenAI limits
        # Typical standard tiers support significantly more than 3 concurrent requests.
        # REDUCED TO 3 to prevent "Missing Findings" due to Rate Limit Drops (429)
        with ThreadPoolExecutor(max_workers=3) as executor:
            future_to_chunk = {
                executor.submit(self._process_chunk, chunk, guidelines): chunk 
                for chunk in chunks
            }
            
            for future in as_completed(future_to_chunk):
                try:
                    chunk_findings = future.result()
                    findings.extend(chunk_findings)
                except Exception as e:
                    logger.error(f"Chunk processing failed: {e}")
        
        return findings

    def _create_chunks(self, model: DocumentModel) -> List[Dict[str, Any]]:
        """
        Create chunks: Aggregates sections to reduce API calls (Smart Chunking).
        Merges sequential sections until a size threshold is reached.
        """
        chunks = []
        
        # Buffers for current chunk
        current_text_buffer = ""
        current_tables = []
        current_para_ids = []
        # We track the 'start' section for the chunk to use as location_id base
        current_location_id = None
        
        # Soft limit for chunk size (characters). 
        # ~4000 chars is roughly 1000 tokens, ideal for speed.
        CHUNK_CHAR_LIMIT = 4000 
        
        for i, section in enumerate(model.sections):
            if not section.content_ids:
                continue
                
            # Content Text
            display_heading = f"## {section.heading_text}\n" if section.heading_level > 0 else ""
            section_text = f"{display_heading}"
            sec_tables = []
            sec_paras = []
            
            for cid in section.content_ids:
                if cid in model.paragraphs:
                    p = model.paragraphs[cid]
                    section_text += f"[ID: {cid}] {p.text}\n"
                    sec_paras.append(cid)
                elif cid in model.tables:
                    t = model.tables[cid]
                    tbl_text = f"[ID: {cid}] [Table {cid}: headers={t.headers}]"
                    if t.content:
                        for r_i, row in enumerate(t.content):
                            cell_items = []
                            for c_i, cell in enumerate(row):
                                cell_id = f"{cid}_r{r_i}_c{c_i}"
                                cell_items.append(f"[ID: {cell_id}] {str(cell or '').strip()}")
                            row_str = " | ".join(cell_items)
                            tbl_text += f"\n  Row {r_i+1}: {row_str}"
                    sec_tables.append(tbl_text)
                    section_text += tbl_text + "\n"

            if not section_text.strip(): continue

            # CHECK OPTIMIZATION: Flush if limit exceeded
            if current_text_buffer and (len(current_text_buffer) + len(section_text) > CHUNK_CHAR_LIMIT):
                chunks.append({
                    "chunk_text": current_text_buffer,
                    "metadata": {
                        "previous_section": "dynamic_merged",
                        "next_section": section.heading_text,
                        "tables": current_tables,
                        "location_id": current_location_id,
                        "para_ids_context": current_para_ids[:3]
                    }
                })
                # Reset buffers
                current_text_buffer = ""
                current_tables = []
                current_para_ids = []
                current_location_id = None

            # Accumulate
            if not current_text_buffer:
                current_location_id = section.id
                
            current_text_buffer += section_text + "\n"
            current_tables.extend(sec_tables)
            current_para_ids.extend(sec_paras)

        # Flush final buffer
        if current_text_buffer:
            chunks.append({
                "chunk_text": current_text_buffer,
                "metadata": {
                    "previous_section": "dynamic_merged",
                    "next_section": "End of Document",
                    "tables": current_tables,
                    "location_id": current_location_id,
                    "para_ids_context": current_para_ids[:3]
                }
            })
            
        return chunks

    def _process_chunk(self, chunk: Dict, guidelines: List[Dict]) -> List[Finding]:
        chunk_findings = []
        
        # --- PRE-PROCESS: HARD DISABLE FILTER ---
        safe_text = llm_safe_text(chunk['chunk_text'])
        if not safe_text.strip():
             return []

        # --- STEP 1: CLASSIFICATION ---
        class_prompt = self.prompt_builder.build_classification_prompt(safe_text)
        class_response = self.llm_client.generate_json(class_prompt, retries=2)
        
        # Default context if classification fails or returns empty
        content_context = {
            "content_types": ["TECHNICAL_CONTENT"], 
            "should_check_grammar": True,
            "should_check_spelling": True,
            "should_check_acronyms": True
        }
        
        if class_response and isinstance(class_response, dict):
            # Merge defaults with response to be safe
            content_context.update(class_response)
        elif class_response and isinstance(class_response, list) and len(class_response) > 0:
             # Handle case if LLM returns a list (unlikely based on prompt, but robust)
             if isinstance(class_response[0], dict):
                  content_context.update(class_response[0])

        content_types = content_context.get("content_types", [])
        print(f"DEBUG: Chunk classification: {content_types}") # DEBUG PRINT
        
        # FORCE FALLBACK: If content_types is just technical, assume it covers REQUIREMENTS too for safety 
        # (This addresses user issue where 'Ambiguous Requirement' is filtered out because classification didn't say 'REQUIREMENTS')
        if "TECHNICAL_CONTENT" in content_types and "REQUIREMENTS" not in content_types:
             content_types.append("REQUIREMENTS")
        
        # --- STEP 2: DYNAMIC GUIDELINE SELECTION ---
        active_rules = []
        
        # Iterate over all available guidelines
        for guideline in guidelines:
            rule_id = guideline.get('id', '')
            
            # Determine category for this rule
            # Try to map Rule ID prefix or known logic to Matrix Keys
            # We use GUIDELINE_CATEGORY_MAP for this.
            
            rule_assigned_categories = []
            
            # 1. Check heuristics on Rule ID
            for cat, prefixes in GUIDELINE_CATEGORY_MAP.items():
                if any(pre in rule_id for pre in prefixes):
                    rule_assigned_categories.append(cat)
            
            # If no category found, default to 'TECHNICAL_CONTENT' applicable or include it?
            # User wants "DO NOT check rules that are not applicable".
            # If we enable "STRUCTURE", it applies to Technical Content usually.
            if not rule_assigned_categories:
                # Default fallback: Treat as generic technical rule
                rule_assigned_categories = ["GRAMMAR", "STRUCTURE"] 

            # Check if ANY of the rule's categories are active for the current content types
            is_active = False
            for rule_cat in rule_assigned_categories:
                # Get allowed content types for this rule category from MATRIX
                allowed_types = GUIDELINE_MATRIX.get(rule_cat, [])
                
                # Intersection: Does current content_types overlap with allowed types?
                if any(t in content_types for t in allowed_types):
                    is_active = True
                    break
            
            if is_active:
                active_rules.append(guideline)
        
        # Logic: If no rules active, skip validation? Or fallback?
        # FORCE INJECT CRITICAL RULES: Ensure ambiguity and context are ALWAYS checked regardless of classification
        # This fixes non-deterministic "missing results" issue.
        critical_ids = {"AMBIGUOUS_REQUIREMENT", "MISSING_CONTEXT", "LOGICAL_CONSISTENCY"}
        for guideline in guidelines:
            if guideline.get('id') in critical_ids:
                # Add if not already present
                if not any(r.get('id') == guideline.get('id') for r in active_rules):
                    active_rules.append(guideline)

        if not active_rules:
             logger.info(f"No active rules for content types: {content_types}")
             return []

        # --- STEP 3: DYNAMIC VALIDATION ---
        prompt = self.prompt_builder.build_dynamic_validation_prompt(
            safe_text, 
            content_context, 
            active_rules
        )
        
        raw_response = self.llm_client.generate_json(prompt, retries=2)
        
        if not raw_response:
             return []
             
        # Normalize list/dict
        if isinstance(raw_response, dict): 
            raw_response = [raw_response]
            
        for item in raw_response:
            try:
                # Validate schema
                if 'issue' not in item or 'rule_id' not in item:
                    continue
                    
                finding = Finding(
                    category=item.get('category', 'Content'),
                    rule_id=item.get('rule_id', 'GPT_ERROR'),
                    issue=item.get('issue', 'Issue identified by AI'),
                    # Map GPT's location format back if needed, but we used IDs from text
                    location_id=item.get('location_id', 'unknown'), 
                    suggestion=item.get('suggestion', ''),
                    evidence=item.get('evidence', ''),
                    confidence=float(item.get('confidence', 0.8))
                )
                
                # Post-processing: Correct location type if needed
                if 'table' in finding.location_id.lower():
                    finding.location_type = 'table'
                else:
                    finding.location_type = 'paragraph'
                    
                chunk_findings.append(finding)
            except Exception as e:
                logger.error(f"Error parsing finding: {e}")
                
        return chunk_findings

    def _verify_finding(self, finding: Dict) -> bool:
        """Second-pass LLM filter"""
        try:
            prompt = self.prompt_builder.build_verification_prompt(finding)
            response = self.llm_client.client.chat.completions.create(
                model=self.llm_client.deployment_name,
                messages=[{"role": "user", "content": prompt}],
                temperature=0,
                max_tokens=10
            )
            ans = response.choices[0].message.content.strip().lower()
            return "true" in ans
        except:
            return True # Fallback to accept if verification fails to run


# ============================================================================
# VALIDATION FUNCTIONS (from validator.py)
# ============================================================================

_CONFIG_PATH = os.path.join(os.path.dirname(__file__), 'data', 'validation_config.json')
try:
    with open(_CONFIG_PATH, 'r') as f:
        DEFAULT_CONFIG = json.load(f)
except Exception:
    DEFAULT_CONFIG = {
        "toc_similarity_threshold": 0.80,
        "min_words_in_section": 8,
        "revision_order": "desc",
        "date_format_regex": r"^(0?[1-9]|1[0-2])/(0?[1-9]|[12]\d|3[01])/\d{4}$",
        "max_url_head_timeout": 3,
        "min_na_justification_words": 5,
        "font_requirements": {
            "body": {"name": "arial", "size": 12.0},
            "h1": {"name": "arial", "size": 24.0},
            "h2": {"name": "arial", "size": 18.0},
            "h3": {"name": "arial", "size": 16.0}
        }
    }

def normalize_text(s: str) -> str:
    return re.sub(r'\s+', ' ', re.sub(r'[^\w\s]', '', (s or "").lower())).strip()

def fuzzy_eq(a: str, b: str, thresh: float = 0.90) -> bool:
    if not a or not b: return False
    return SequenceMatcher(None, normalize_text(a), normalize_text(b)).ratio() >= thresh

def extract_urls(text: str) -> List[str]:
    # Changed regex to catch trailing schemes like https:// (common error)
    raw_urls = re.findall(r'(?:https?://|ftp://|www\.)[^\s<>"{}|\\^`\[\]]*', text)
    clean_urls = []
    for u in raw_urls:
        clean = u.rstrip('.,;:!)?')
        if clean:
            clean_urls.append(clean)
    return list(dict.fromkeys(clean_urls))

def is_valid_url_format(url: str) -> bool:
    """Validate URL format using regex"""
    url_pattern = re.compile(
        r'^(?:http|ftp)s?://'  # http:// or https://
        r'(?:(?:[A-Z0-9](?:[A-Z0-9-]{0,61}[A-Z0-9])?\.)+(?:[A-Z]{2,6}\.?|[A-Z0-9-]{2,}\.?)|'  # domain
        r'localhost|'  # localhost
        r'\d{1,3}\.\d{1,3}\.\d{1,3}\.\d{1,3})'  # or IP
        r'(?::\d+)?'  # optional port
        r'(?:/?|[/?]\S+)$', re.IGNORECASE)
    
    # Handle www. URLs
    if url.startswith('www.'):
        url = 'http://' + url
    
    return bool(url_pattern.match(url))

def check_url_with_details(url: str, timeout: int = 3) -> Tuple[Optional[int], Optional[str]]:
    """
    Check URL accessibility with detailed error reporting
    Returns: (status_code, error_type)
    error_type can be: None, 'DNS_FAILURE', 'TIMEOUT', 'CONNECTION_ERROR', 'INVALID_URL'
    """
    if url.startswith('www.'):
        url = 'http://' + url
    
    try:
        r = requests.head(url, timeout=timeout, allow_redirects=True)
        return r.status_code, None
    except requests.exceptions.Timeout:
        return None, 'TIMEOUT'
    except requests.exceptions.ConnectionError as e:
        if 'Name or service not known' in str(e) or 'getaddrinfo failed' in str(e):
            return None, 'DNS_FAILURE'
        return None, 'CONNECTION_ERROR'
    except requests.exceptions.InvalidURL:
        return None, 'INVALID_URL'
    except requests.RequestException:
        # Try GET as fallback
        try:
            r = requests.get(url, timeout=timeout, allow_redirects=True)
            return r.status_code, None
        except requests.exceptions.Timeout:
            return None, 'TIMEOUT'
        except requests.exceptions.ConnectionError as e:
            if 'Name or service not known' in str(e) or 'getaddrinfo failed' in str(e):
                return None, 'CONNECTION_ERROR'
            return None, 'CONNECTION_ERROR'
        except Exception:
            return None, 'CONNECTION_ERROR'

def check_url_head_get(url: str, timeout: int = 3) -> Optional[int]:
    """Legacy function for backward compatibility"""
    status, _ = check_url_with_details(url, timeout)
    return status

def parse_date_strict(date_str: str, format_hint: str = None) -> Optional[datetime]:
    """
    Parse date with strict validation
    format_hint: 'MM/DD/YYYY', 'DD-MMM-YYYY', 'YYYY-MM-DD', etc.
    """
    format_map = {
        'MM/DD/YYYY': '%m/%d/%Y',
        'DD-MMM-YYYY': '%d-%b-%Y',
        'YYYY-MM-DD': '%Y-%m-%d',
        'DD/MM/YYYY': '%d/%m/%Y',
        'MM-DD-YYYY': '%m-%d-%Y'
    }
    
    if format_hint and format_hint in format_map:
        try:
            return datetime.strptime(date_str, format_map[format_hint])
        except ValueError:
            return None
    
    # Try all formats
    for fmt in format_map.values():
        try:
            return datetime.strptime(date_str, fmt)
        except ValueError:
            continue
    
    # Fallback to dateutil
    try:
        return date_parser.parse(date_str, dayfirst=False, fuzzy=False)
    except Exception:
        return None


def find_section_by_heading(model: DocumentModel, name_terms: List[str]) -> Optional[SectionModel]:
    lower_terms = [t.lower() for t in name_terms]
    for s in model.sections:
        if not s.heading_text: continue
        ht = s.heading_text.lower().strip()
        for t in lower_terms:
            if ht == t or t in ht or fuzzy_eq(ht, t, 0.9):
                return s
    return None

def gather_all_text(model: DocumentModel) -> str:
    paras = sorted(model.paragraphs.values(), key=lambda p: p.index)
    return "\n".join([p.text or "" for p in paras])

def word_count(s: str) -> int:
    return len([w for w in re.findall(r'\w+', s) if w.strip()])

def parse_date_try(date_str: str) -> Optional[datetime]:
    try:
        dt = date_parser.parse(date_str, dayfirst=False, fuzzy=True)
        return dt
    except Exception:
        return None

def _extract_font_requirements(template_model: DocumentModel) -> Dict[str, Dict[str, Any]]:
    """
    Extracts font name and size requirements from the template document using a majority vote.
    Returns: {
        "body": {"name": str, "size": float},
        "h1": {"name": str, "size": float},
        "h2": {"name": str, "size": float},
        "h3": {"name": str, "size": float}
    }
    """
    if not template_model:
        return {}

    # stats[level] = {"names": {name: count}, "sizes": {size: count}}
    stats = defaultdict(lambda: {"names": defaultdict(int), "sizes": defaultdict(int)})

    for para in template_model.paragraphs.values():
        if not para.text.strip():
            continue
        
        # Determine role
        level = para.heading_level
        if level > 3: level = 3 # Group all sub-sections as h3+
        
        if para.font_name:
            stats[level]["names"][para.font_name] += 1
        if para.font_size > 0:
            stats[level]["sizes"][para.font_size] += 1

    requirements = {}
    roles = {0: "body", 1: "h1", 2: "h2", 3: "h3"}

    for level, role in roles.items():
        role_stats = stats.get(level)
        if role_stats and role_stats["names"] and role_stats["sizes"]:
            # Pick majority
            best_name = max(role_stats["names"].items(), key=lambda x: x[1])[0]
            best_size = max(role_stats["sizes"].items(), key=lambda x: x[1])[0]
            requirements[role] = {"name": best_name, "size": float(best_size)}
    
    return requirements

def parse_version_tuple(v: str) -> Tuple:
    nums = re.findall(r'\d+', v)
    return tuple(int(x) for x in nums) if nums else ()

# ============================================================================
# FIXED VALIDATION LOGIC
# ============================================================================

def iter_all_text_elements(model: DocumentModel):
    """Helper to iterate over all text checking locations (paragraphs and table cells)"""
    for pid, p in model.paragraphs.items():
        if '_r' in pid and '_c' in pid:
            yield pid, "table", p.text or ""
        else:
            yield pid, "paragraph", p.text or ""

def get_last_paragraph_id(model: DocumentModel) -> str:
    """Helper to find the last paragraph ID for appending comments at the end"""
    if not model.paragraphs:
        return "document"
    # Find paragraph with highest global index
    last_para = max(model.paragraphs.values(), key=lambda p: p.global_index)
    return last_para.id

def validate_acronym_section_presence(model: DocumentModel, template_model: DocumentModel = None) -> List[Finding]:
    """Check for presence of Acronyms/Definitions section. Anchor to end if missing."""
    # Check input document
    for section in model.sections:
        title = (section.heading_text or "").lower()
        if any(term in title for term in ['acronym', 'definition', 'abbreviation']):
            return []
            
    loc_id = get_last_paragraph_id(model)
    return [Finding(
        category="Structure", 
        rule_id="MISSING_ACRONYM_SECTION", 
        issue="Document is missing an 'Acronyms and Definitions' section.", 
        location_id=loc_id, 
        suggestion="Add an 'Acronyms and Definitions' section at the end of the document.",
        priority="MEDIUM"
    )]

def validate_toc_presence(input_model: DocumentModel, template_model: DocumentModel = None) -> List[Finding]:
    """
    Checks for TOC presence in the first 2 pages (~20 paragraphs).
    """
    findings = []
    
    # Check first 20 paragraphs (approx 2 pages)
    toc_found_in_range = False
    sorted_paras = sorted(input_model.paragraphs.items(), key=lambda x: x[1].index)
    
    # Scan first 20 elements for a TOC header
    for _, para in sorted_paras[:20]:
        t = (para.text or "").lower().strip()
        if t in ['table of contents', 'contents', 'toc']:
            toc_found_in_range = True
            break
            
    if not toc_found_in_range:
        # Check if it exists ELSEWHERE (fallback)
        toc_section = find_section_by_heading(input_model, ['table of contents', 'contents', 'toc'])
        if toc_section:
            # It exists but NOT in the first 2 pages
            findings.append(Finding(
                category="Structure", 
                rule_id="TOC_PLACEMENT_ISSUE", 
                issue="Table of Contents (TOC) is not placed within the first 2 pages.", 
                location_id=toc_section.id, 
                location_type="section",
                suggestion="Move the Table of Contents to the beginning of the document (page 2)."
            ))
        else:
            loc_id = get_last_paragraph_id(input_model)
            findings.append(Finding(
                category="Structure", 
                rule_id="MISSING_TOC", 
                issue="Table of Contents (TOC) not found in the first 2 pages.", 
                location_id=loc_id, 
                suggestion="Please add a Table of Contents after the title page (usually page 2)."
            ))
    else:
        # Check if empty
        toc_section = find_section_by_heading(input_model, ['table of contents', 'contents', 'toc'])
        if toc_section:
            toc_text = ""
            for cid in toc_section.content_ids:
                if cid in input_model.paragraphs:
                    toc_text += input_model.paragraphs[cid].text or ""
            if word_count(toc_text) < 3:
                findings.append(Finding(
                    category="Structure", 
                    rule_id="EMPTY_TOC", 
                    issue="The Table of Contents appears to be empty.", 
                    location_id=toc_section.id, 
                    suggestion="Ensure the TOC is updated and reflects the document structure."
                ))
                
    return findings

def validate_document_similarity(input_model: DocumentModel, template_model: DocumentModel, config=DEFAULT_CONFIG) -> List[Finding]:
    """
    Consolidated TOC similarity validation.
    Redirects to DocumentProcessor._check_toc_similarity to ensure consistent anchoring logic.
    """
    if not template_model:
        return []
        
    dp = DocumentProcessor()
    # Ensure any config overrides are passed if necessary
    similarity_score, findings = dp._check_toc_similarity(input_model, template_model)
    
    if similarity_score < 90.0:
        findings.append(Finding(
            category="Structure", 
            rule_id="LOW_TOC_SIMILARITY", 
            issue=f"Document structure similarity is low ({similarity_score:.1f}%). Target is 90%.", 
            location_id="document", 
            suggestion="Align document sections more closely with the template."
        ))
    return findings


def validate_hyperlinks(model: DocumentModel, config=DEFAULT_CONFIG) -> List[Finding]:
    """Enhanced URL validation checking paragraphs AND tables"""
    findings = []
    
    # Check text URLs in paragraphs and tables
    for loc_id, loc_type, text in iter_all_text_elements(model):
        if not text: continue
        urls = extract_urls(text)
        for url in urls:
            if not is_valid_url_format(url):
                findings.append(Finding(category="Content", rule_id="MALFORMED_URL", issue=f"Malformed URL: {url}", location_id=loc_id, location_type=loc_type, suggestion="Fix URL format"))
                continue
            
            status, error_type = check_url_with_details(url, timeout=config.get("max_url_head_timeout", 3))
            if status and status >= 400:
                findings.append(Finding(category="Content", rule_id="BROKEN_URL", issue=f"Broken URL (HTTP {status}): {url}", location_id=loc_id, location_type=loc_type, suggestion="Fix broken link"))
            elif error_type:
                 findings.append(Finding(category="Content", rule_id="BROKEN_URL", issue=f"Unreachable URL ({error_type}): {url}", location_id=loc_id, location_type=loc_type, suggestion="Verify URL"))

    # Check embedded hyperlinks (from relationships)
    for url in set(model.hyperlinks):
         if not is_valid_url_format(url):
             findings.append(Finding(category="Content", rule_id="MALFORMED_HYPERLINK", issue=f"Malformed embedded link: {url}", location_id="document", suggestion="Fix hyperlink target format"))
             continue
             
         status, error_type = check_url_with_details(url, timeout=config.get("max_url_head_timeout", 3))
         if status and status >= 400:
             findings.append(Finding(category="Content", rule_id="BROKEN_HYPERLINK", issue=f"Broken embedded link (HTTP {status}): {url}", location_id="document", suggestion="Fix broken hyperlink target"))
         elif error_type:
             findings.append(Finding(category="Content", rule_id="BROKEN_HYPERLINK", issue=f"Broken embedded link ({error_type}): {url}", location_id="document", suggestion="Verify hyperlink target"))

    return findings

def validate_references(model: DocumentModel, template_model: DocumentModel = None) -> List[Finding]:
    """Check for presence of References section. Anchor to end if missing."""
    # Check input document
    for section in model.sections:
        title = (section.heading_text or "").lower()
        if 'reference' in title:
            return []
            
    loc_id = get_last_paragraph_id(model)
    return [Finding(
        category="Content", 
        rule_id="MISSING_REFERENCES", 
        issue="Document is missing a 'References' section.", 
        location_id=loc_id, 
        suggestion="Please add a References section toward the end of the document.",
        priority="MEDIUM"
    )]


def validate_fonts(model: DocumentModel, config=DEFAULT_CONFIG, font_requirements: Dict[str, Any] = None) -> List[Finding]:
    findings = []
    
    # Use template requirements if provided, else fallback to config
    base_reqs = config.get('font_requirements', {})
    if font_requirements:
        # Merge or override? Let's override what's in template, keep others.
        final_reqs = {**base_reqs, **font_requirements}
        source_label = "Template"
    else:
        final_reqs = base_reqs
        source_label = "GDP"

    for para_id, para in model.paragraphs.items():
        if not para.text.strip(): continue
        
        # Determine Role
        role = "body"
        role_display = "Body Content"
        if para.heading_level == 1:
            role = "h1"
            role_display = "Main Header (Heading 1)"
        elif para.heading_level == 2:
            role = "h2"
            role_display = "Section header (Heading 2)"
        elif para.heading_level >= 3:
            role = "h3"
            role_display = f"Sub section header (Heading {para.heading_level})"

        req = final_reqs.get(role)
        if not req: continue # Should not happen with defaults

        target_size = req.get("size", 12.0)
        
        # 1. Font Size Consistency
        if para.font_size > 0:
            if abs(para.font_size - target_size) > 0.5:
                # Use standard suggestion wording from plan
                if para.heading_level == 1:
                    suggestion = f"As per {source_label} standards, Main Headers (Heading 1) must be {target_size}pt. Please update the font size."
                elif para.heading_level == 2:
                    suggestion = f"As per {source_label} standards, Section headers (Heading 2) must be {target_size}pt. Please update the font size."
                elif para.heading_level >= 3:
                    suggestion = f"As per {source_label} standards, Sub section headers (Heading {para.heading_level}) must be {target_size}pt. Please update the font size."
                else:
                    suggestion = f"As per {source_label} standards, Body Content must be {target_size}pt. Please update the font size."
                
                findings.append(Finding(
                    category="Formatting", 
                    rule_id="INCORRECT_FONT_SIZE", 
                    issue=f"{role_display} uses {para.font_size}pt font size.", 
                    location_id=para_id, 
                    suggestion=suggestion,
                    evidence=para.text[:50] if len(para.text) > 50 else para.text
                ))
        
        # 2. Heading Specific Consistency (No Period)
        if para.heading_level > 0:
            if para.text.strip().endswith('.'):
                last_dot_idx = para.text.rfind('.')
                findings.append(Finding(
                    category="Formatting",
                    rule_id="HEADING_WITH_PERIOD",
                    issue="Heading ends with a period.",
                    location_id=para_id,
                    suggestion=f"As per {source_label} standards, headers should not end with a period. Please remove the period.",
                    evidence=".",
                    evidence_span=(last_dot_idx, last_dot_idx + 1)
                ))
            if para.heading_level == 1:
                # Title Case check
                words = [w for w in para.text.split() if w[0].isalpha()]
                if words:
                    minor_words = {'a', 'an', 'the', 'and', 'but', 'or', 'for', 'nor', 'on', 'at', 'to', 'from', 'by', 'of', 'in', 'with'}
                    not_capitalized = [w for w in words if w[0].islower() and w.lower() not in minor_words]
                    if not_capitalized:
                        findings.append(Finding(
                            category="Formatting",
                            rule_id="HEADING_CASE_ISSUE",
                            issue="Main Header is not in Title Case.",
                            location_id=para_id,
                            suggestion=f"As per {source_label} standards, Main Headers should follow Title Case."
                        ))
            
            # Heading Boldness Check
            if not para.is_bold:
                 findings.append(Finding(
                    category="Formatting",
                    rule_id="HEADING_NOT_BOLD",
                    issue=f"Heading '{para.text[:30]}...' is not bold.",
                    location_id=para_id,
                    suggestion=f"As per {source_label} standards, all headings must be Bold."
                ))

        # 3. Bullet & Sub-bullet Consistency (Trailing Period)
        is_bullet = any(x in para.style.lower() for x in ['bullet', 'list', 'listparagraph']) or (para.numbering is not None)
        if is_bullet and para.text.strip():
            if not para.text.strip().endswith('.'):
                findings.append(Finding(
                    category="Formatting", 
                    rule_id="MISSING_PERIOD_BULLET", 
                    issue="Bullet/sub-bullet point is missing a trailing period (.).", 
                    location_id=para_id, 
                    suggestion=f"As per {source_label} consistency guidelines, please add a period at the end of each bullet point.",
                    evidence=para.text[-10:] if len(para.text) > 10 else para.text
                ))

    return findings

def validate_whitespace(model: DocumentModel) -> List[Finding]:
    findings = []
    for loc_id, loc_type, text in iter_all_text_elements(model):
        if not text: continue
        
        if "  " in text:
            findings.append(Finding(
                category="Formatting", 
                rule_id="MULTIPLE_SPACES", 
                issue="Multiple spaces found", 
                location_id=loc_id, 
                location_type=loc_type,
                suggestion="Use single space"
            ))
            
        for line in text.split('\n'):
            if line != line.rstrip():
                findings.append(Finding(
                    category="Formatting", 
                    rule_id="TRAILING_SPACES", 
                    issue="Trailing spaces found", 
                    location_id=loc_id, 
                    location_type=loc_type,
                    suggestion="Remove trailing spaces"
                ))
                break
    return findings

def validate_date_format_and_values(model: DocumentModel, config=DEFAULT_CONFIG) -> List[Finding]:
    """Enhanced date validation - enforces strict DD/MM/YYYY"""
    findings = []
    
    # Strictly check for slashes. We flag other separators as invalid format.
    invalid_separator_patterns = {
        'DD-MMM-YYYY': r'\b(0?[1-9]|[12]\d|3[01])-(Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)-\d{4}\b',
        'YYYY-MM-DD': r'\b\d{4}-(0?[1-9]|1[0-2])-(0?[1-9]|[12]\d|3[01])\b',
        'MM-DD-YYYY': r'\b(0?[1-9]|1[0-2])-(0?[1-9]|[12]\d|3[01])-\d{4}\b',
        'DD.MM.YYYY': r'\b(0?[1-9]|[12]\d|3[01])\.(0?[1-9]|1[0-2])\.\d{4}\b',
        # New: Catch MMM/DD/YY or MMM/DD/YYYY (e.g. OCT/29/25) which is often missed because of alpha
        'MMM/DD/YYYY': r'\b(Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)(?:uary|ruary|ch|il|ust|ember|ober)?/\d{1,2}/\d{2,4}\b'
    }
    
    for loc_id, loc_type, text in iter_all_text_elements(model):
        if not text: continue
        
        # 1. Check for hyphenated, dotted, or alpha-slash dates (Invalid Separators/Format)
        for format_name, pattern in invalid_separator_patterns.items():
            for match in re.finditer(pattern, text, re.IGNORECASE):
                date_str = match.group()
                findings.append(Finding(
                    category="Content", 
                    rule_id="INVALID_DATE_FORMAT", 
                    issue=f"Invalid date format '{date_str}' ({format_name}).", 
                    location_id=loc_id, 
                    location_type=loc_type, 
                    suggestion="Use slash format (DD/MM/YYYY) for GDP compliance."
                ))

        # 2. Check for slash-separated dates (Validate logic and bounds)
        slash_matches = re.finditer(r'\b(\d{1,2})/(\d{1,2})/(\d{2,4})\b', text)
        
        for match in slash_matches:
            date_str = match.group()
            p1, p2, year_str = match.group(1), match.group(2), match.group(3)
            p1_int, p2_int = int(p1), int(p2)
            
            # Rule A: 4-digit Year
            if len(year_str) == 2:
                 findings.append(Finding(
                     category="Content", 
                     rule_id="INVALID_DATE_FORMAT", 
                     issue=f"Date '{date_str}' uses a 2-digit year.", 
                     location_id=loc_id, 
                     location_type=loc_type, 
                     suggestion="Use a 4-digit year (e.g., 2025)."
                 ))
                 continue
            
            p3_int = int(year_str)
            
            # Rule B: Strict DD/MM/YYYY Validation
            # DD must be 1-31, MM must be 1-12
            is_valid_dd_mm = (1 <= p1_int <= 31) and (1 <= p2_int <= 12)
            is_valid_mm_dd = (1 <= p1_int <= 12) and (1 <= p2_int <= 31)
            
            if is_valid_dd_mm:
                try:
                    # Check if actual date is logical (e.g. Feb 30)
                    datetime(p3_int, p2_int, p1_int)
                except ValueError:
                    findings.append(Finding(
                        category="Content", 
                        rule_id="INVALID_DATE_VALUE", 
                        issue=f"'{date_str}' is not a valid calendar date.", 
                        location_id=loc_id, 
                        location_type=loc_type, 
                        suggestion="Provide a valid date in DD/MM/YYYY format."
                    ))
                    continue

                # Context check: If it COULD be MM/DD/YYYY and DD is <= 12, we can't be sure, 
                # but if DD > 12 and MM <= 12, it's definitely DD/MM/YYYY.
                # If MM > 12 and DD <= 12, it's definitely MM/DD/YYYY (Invalid).
                if p1_int > 12 and p2_int <= 12:
                    # Confirmed DD/MM/YYYY
                    pass 
                elif p1_int <= 12 and p2_int > 12:
                    # Confirmed MM/DD/YYYY - Violation
                    findings.append(Finding(
                        category="Content", 
                        rule_id="INVALID_DATE_FORMAT", 
                        issue=f"Date '{date_str}' appears to be in MM/DD/YYYY (US) format.", 
                        location_id=loc_id, 
                        location_type=loc_type, 
                        suggestion="As per GDP standards, use strict DD/MM/YYYY format (Day first)."
                    ))
            elif is_valid_mm_dd:
                 # Definitely MM/DD/YYYY since is_valid_dd_mm failed (so p2_int > 12 or p1_int > 31)
                 findings.append(Finding(
                     category="Content", 
                     rule_id="INVALID_DATE_FORMAT", 
                     issue=f"Date '{date_str}' is in MM/DD/YYYY (US) format.", 
                     location_id=loc_id, 
                     location_type=loc_type, 
                     suggestion="Use strict DD/MM/YYYY format."
                 ))
            else:
                 # Neither valid
                 findings.append(Finding(
                     category="Content", 
                     rule_id="INVALID_DATE_VALUE", 
                     issue=f"Invalid date numbers in '{date_str}'.", 
                     location_id=loc_id, 
                     location_type=loc_type, 
                     suggestion="Ensure Day (1-31) and Month (1-12) are correct (DD/MM/YYYY)."
                 ))

    return findings

def validate_revision_date_order(model: DocumentModel) -> List[Finding]:
    """Ensure revision history dates are in chronological order (newest first)"""
    findings = []
    print("DEBUG: Running validate_revision_date_order")
    for table_id, table in model.tables.items():
        if not table.headers: continue
        headers = [h.lower().strip() for h in table.headers]
        
        date_idx = next((i for i, h in enumerate(headers) if 'date' in h or 'revision' in h), None)
        if date_idx is None: continue
        
        dates = []
        for row_idx, row in enumerate(table.content[1:], 1):
            if len(row) <= date_idx: continue
            date_str = row[date_idx].strip()
            if not date_str or date_str.lower() in ['n/a', 'na', '-', '']: continue
            
            # Try parsing leniently for history checks
            parsed = parse_date_try(date_str)
            if parsed:
                dates.append((row_idx, parsed, date_str))
        
        if len(dates) >= 2:
            for i in range(len(dates) - 1):
                # Check Dates: Newest should be first (i.e. date[i] > date[i+1])
                # Allow equal dates
                if dates[i][1] < dates[i+1][1]: 
                    findings.append(Finding(
                        category="Authorship", 
                        rule_id="REVISION_DATE_ORDER", 
                        issue=f"Date order incorrect: {dates[i][2]} is before {dates[i+1][2]}", 
                        location_id=table_id, 
                        location_type="table", 
                        suggestion="Order newest to oldest",
                        priority="HIGH"
                    ))
                    break
        
        # Also check version order if version column exists
        version_idx = next((i for i, h in enumerate(headers) if 'ver' in h or 'rev' in h), None)
        if version_idx is not None:
             # Extract versions
             versions = []
             for row_idx, row in enumerate(table.content[1:], 1):
                 if len(row) <= version_idx: continue
                 v_str = row[version_idx].strip()
                 if v_str:
                     versions.append((row_idx, v_str))
             
             if len(versions) >= 2:
                 for i in range(len(versions) - 1):
                     v1_str = versions[i][1]
                     v2_str = versions[i+1][1]
                     # Try to parse simple X.Y versioning
                     try:
                         # Normalize 1.0 to [1, 0]
                         p1 = [int(p) for p in re.findall(r'\d+', v1_str)]
                         p2 = [int(p) for p in re.findall(r'\d+', v2_str)]
                         if p1 and p2:
                             # Should be descending: p1 >= p2
                             if p1 < p2:
                                 findings.append(Finding(
                                     category="Authorship", 
                                     rule_id="VERSION_ORDER", 
                                     issue=f"Version order incorrect: {v1_str} appears before {v2_str}", 
                                     location_id=table_id, 
                                     location_type="table", 
                                     suggestion="Order newest version (highest number) to oldest",
                                     priority="MEDIUM"
                                 ))
                                 break
                     except:
                         pass

    return findings


def validate_version_consistency(model: DocumentModel) -> List[Finding]:
    findings = []
    full_text = gather_all_text(model)
    # Improved regex for finding "Version 1.0" or "v1.0" in headers
    header_versions = set(re.findall(r'(?:Version|v)\s*[:\.]?\s*(\d+(?:\.\d+)+)', full_text, flags=re.I))
    
    table_versions = []
    revision_table_id = None
    
    for table_id, table in model.tables.items():
        headers = [h.lower() for h in (table.headers or [])]
        vidx = next((i for i,h in enumerate(headers) if 'version' in h or 'rev' in h), None)
        if vidx is not None:
            revision_table_id = table_id
            for row in table.content[1:]:
                if vidx < len(row):
                    v_match = re.search(r'(\d+(?:\.\d+)+)', str(row[vidx]))
                    if v_match:
                        table_versions.append(v_match.group(1))

    # Version Order Check (Descending)
    if len(table_versions) >= 2:
        for i in range(len(table_versions) - 1):
             try:
                 v1 = tuple(map(int, table_versions[i].split('.')))
                 v2 = tuple(map(int, table_versions[i+1].split('.')))
                 if v1 < v2:
                     findings.append(Finding(category="Authorship", rule_id="VERSION_ORDER_INVALID", issue=f"Version order invalid: {table_versions[i]} before {table_versions[i+1]}", location_id=revision_table_id or "document", location_type="table", suggestion="Order newest to oldest"))
                     break
             except: pass

    # Mismatch check
    if header_versions and table_versions:
        # Latest form table should match at least one header version
        latest_table = sorted(table_versions, key=lambda v: tuple(map(int, v.split('.'))), reverse=True)[0]
        if latest_table not in header_versions:
             # It might be normal if header has no version, but if it DOES have valid versions and none match latest...
             pass 

    return findings

# --- Legacy dedupe REMOVED: Using dedupe_findings instead ---

# Functions already defined in FIXED VALIDATION LOGIC above.
# validate_toc_presence (line 2193)
# validate_empty_cells, validate_empty_sections, etc. will be moved up if missing.

def validate_empty_cells(model: DocumentModel) -> List[Finding]:
    findings = []
    for table_id, table in model.tables.items():
        for r_idx, row in enumerate(table.content):
            for c_idx, cell in enumerate(row):
                if not (str(cell or "").strip()):
                    findings.append(Finding(
                        category="Structure", 
                        rule_id="EMPTY_CELLS", 
                        issue=f"Empty cell at row {r_idx+1}, col {c_idx+1}", 
                        location_id=f"{table_id}_r{r_idx}_c{c_idx}", 
                        location_type="table", 
                        suggestion="Provide the required information in this cell"
                    ))
    return findings

def validate_empty_sections(model: DocumentModel, config=DEFAULT_CONFIG) -> List[Finding]:
    findings = []
    min_words = config.get("min_words_in_section", 8)
    for section in model.sections:
        if section.heading_level <= 0: continue
        if len(section.content_ids) <= 1:
             loc_id = section.content_ids[0] if section.content_ids else section.id
             findings.append(Finding(category="Structure", rule_id="EMPTY_SECTION", issue=f"Section '{section.heading_text}' appears empty or too brief", location_id=loc_id, location_type="section", suggestion="Provide detailed content for this section", priority="HIGH"))
             continue
        text = ""
        for cid in section.content_ids:
            if cid in model.paragraphs: text += " " + (model.paragraphs[cid].text or "")
            elif cid in model.tables:
                t = model.tables[cid]
                for r in t.content: text += " " + " ".join([c for c in r if c])
        if not text.strip() or word_count(text) < min_words:
            loc_id = section.content_ids[0] if section.content_ids else section.id
            findings.append(Finding(category="Structure", rule_id="EMPTY_SECTION", issue=f"Section '{section.heading_text}' appears empty or too brief", location_id=loc_id, location_type="section", suggestion="Provide detailed content", priority="HIGH"))
    return findings

def validate_spelling(model: DocumentModel, known_acronyms: Set[str] = None, config: Dict[str, Any] = None) -> List[Finding]:
    if config is None: config = DEFAULT_CONFIG
    findings = []
    spell = get_spellchecker()
    ignore_words = set(known_acronyms or [])
    try:
        am = AcronymManager()
        ignore_words.update(am.corporate_acronyms.keys()); ignore_words.update(am.common_acronyms); ignore_words.update(am.pharma_acronyms); ignore_words.update(am.whitelist)
    except: pass
    ignore_words.update(config.get('spelling_ignore_words', []))
    for loc_id, loc_type, text in iter_all_text_elements(model):
        if not text: continue
        text_clean = re.sub(r'https?://\S+|www\.\S+', '', text)
        words = re.findall(r"\b[a-zA-Z']+\b", text_clean)
        for word in words:
            word_clean = word.strip("'")
            if not word_clean or word_clean.lower() in spell or len(word_clean) < 3 or any(c.isdigit() for c in word_clean) or word_clean.isupper(): continue
            if word_clean.lower() in (k.lower() for k in ignore_words): continue
            suggestion = spell.correction(word_clean)
            findings.append(Finding(category="Spelling", rule_id="MISSPELLED_WORD", issue=f"Potential misspelling: '{word_clean}'", location_id=loc_id, location_type=loc_type, suggestion=f"Consider: {suggestion}" if suggestion else "Check spelling", priority="LOW"))
    return findings

def validate_grammar(model: DocumentModel) -> List[Finding]:
    findings = []
    checker = get_grammar_checker()
    if not checker.tool: return findings
    for loc_id, loc_type, text in iter_all_text_elements(model):
        if not text or len(text) < 10: continue
        if '.' not in text and len(text) < 50: continue
        findings.extend(checker.check_text(text, loc_id))
    return findings

def validate_tbd_placeholders(model: DocumentModel) -> List[Finding]:
    findings = []
    patterns = [r'\bTBD\b', r'\bTo be (?:determined|updated|confirmed)\b', r'\bPending\b', r'\bTBU\b', r'\bTBC\b', r'\bTODO\b', r'\bFIXME\b']
    combined = re.compile("|".join(patterns), flags=re.I)
    for loc_id, loc_type, text in iter_all_text_elements(model):
         if not text: continue
         match = combined.search(text)
         if match:
             findings.append(Finding(
                 category="Content", 
                 rule_id="TBD_PLACEHOLDER", 
                 issue=f"Placeholder found: '{match.group()}'", 
                 location_id=loc_id, 
                 location_type=loc_type, 
                 suggestion="Replace placeholder with final content",
                 priority="HIGH"
             ))
    return findings

def validate_author_roles(model: DocumentModel) -> List[Finding]:
    findings = []
    for table_id, table in model.tables.items():
        if not table.headers or len(table.headers) < 2: continue
        headers = [h.lower().strip() for h in table.headers]
        author_idx = next((i for i, h in enumerate(headers) if 'author' in h or 'prepared' in h), None)
        reviewer_idx = next((i for i, h in enumerate(headers) if 'review' in h), None)
        approver_idx = next((i for i, h in enumerate(headers) if 'approv' in h), None)
        if author_idx is None: continue
        for row in table.content[1:]:
            if len(row) <= author_idx: continue
            author = row[author_idx].strip().lower()
            if not author or author in ['n/a', 'na', '-', '']: continue
            if reviewer_idx is not None and len(row) > reviewer_idx and author == row[reviewer_idx].strip().lower():
                findings.append(Finding(category="Authorship", rule_id="AUTHOR_AS_REVIEWER", issue=f"'{row[author_idx].strip()}' is both author and reviewer", location_id=table_id, location_type="table", suggestion="Author should not review their own work"))
            if approver_idx is not None and len(row) > approver_idx and author == row[approver_idx].strip().lower():
                findings.append(Finding(category="Authorship", rule_id="AUTHOR_AS_APPROVER", issue=f"'{row[author_idx].strip()}' is both author and approver", location_id=table_id, location_type="table", suggestion="Author should not approve their own work"))
    return findings

def validate_personal_references(model: DocumentModel) -> List[Finding]:
    findings = []
    personal_pronouns = re.compile(r'\b(I|we|you|me|us|my|our|your|myself|ourselves|yourself|yourselves)\b', re.IGNORECASE)
    for loc_id, loc_type, text in iter_all_text_elements(model):
        if not text: continue
        matches = list(personal_pronouns.finditer(text))
        if matches:
            pronouns_found = list(set(m.group() for m in matches))
            findings.append(Finding(category="Style", rule_id="PERSONAL_REFERENCE" if loc_type == "paragraph" else "PERSONAL_REFERENCE_TABLE", issue=f"Personal pronoun(s) found: {', '.join(pronouns_found[:3])}", location_id=loc_id, location_type=loc_type, suggestion="Use third person or passive voice"))
    return findings

def validate_acronym_first_use(model: DocumentModel) -> List[Finding]:
    findings = []
    try: am = AcronymManager()
    except: return findings
    seen_acronyms = {}; defined_acronyms = set()
    definition_pattern = re.compile(r'([A-Z][A-Za-z\s&-]+)\s*[\(\-]\s*([A-Z]{2,})\s*[\)]?')
    sorted_paras = sorted(model.paragraphs.items(), key=lambda x: x[1].index)
    for para_id, para in sorted_paras:
        if not para.text: continue
        for match in definition_pattern.finditer(para.text): defined_acronyms.add(match.group(2))
        acronyms = am.find_acronyms(para.text)
        for acronym in acronyms:
            base = acronym[:-1] if acronym.endswith('s') and len(acronym) > 2 and acronym[:-1].isupper() else acronym
            if am.is_known_acronym(base) or base in defined_acronyms or len(base) < 2 or base.upper() in {'AM', 'PM', 'ID', 'NO', 'US', 'UK', 'EU'}: continue
            if para.style.lower().startswith('heading') or (para.is_bold and para.font_size >= 12): continue
            if len(acronym) > 10 and not am.is_known_acronym(acronym): continue
            if acronym not in seen_acronyms:
                seen_acronyms[acronym] = para_id
                findings.append(Finding(category="Content", rule_id="UNDEFINED_ACRONYM", issue=f"Acronym '{acronym}' used without definition", location_id=para_id, suggestion=f"Define as 'Full Name ({acronym})' on first use"))
    return findings

def validate_table_headers_consistency(model: DocumentModel, template_model: Optional[DocumentModel] = None) -> List[Finding]:
    findings = []
    if not template_model: return []
    input_tbls = sorted(model.tables.items(), key=lambda x: x[1].index)
    tpl_tbls = sorted(template_model.tables.items(), key=lambda x: x[1].index)
    for i in range(min(len(input_tbls), len(tpl_tbls))):
        i_id, i_tbl = input_tbls[i]; _, t_tbl = tpl_tbls[i]
        if i_tbl.headers != t_tbl.headers:
             findings.append(Finding(category="Structure", rule_id="TABLE_HEADER_MISMATCH", issue=f"Table {i+1} headers mismatch template", location_id=i_id, location_type="table", suggestion=f"Use headers: {', '.join(t_tbl.headers)}"))
    return findings

def run_all_validators(model: DocumentModel, template_model: DocumentModel = None, config: Dict[str, Any] = None, known_acronyms: Set[str]=None) -> List[Finding]:
    """Runs all static validators in parallel and returns aggregated findings."""
    if config is None: config = DEFAULT_CONFIG
    findings = []
    
    # Pre-gather facts
    known_acronyms = set() # Optional
    
    # Extract font requirements from template if available
    font_requirements = _extract_font_requirements(template_model) if template_model else {}
    if font_requirements:
        logger.info(f"Extracted font requirements from template: {font_requirements}")

    static_validators = [
        (validate_toc_presence, (model, template_model)),
        (validate_hyperlinks, (model, config)),
        (validate_empty_cells, (model,)),
        (validate_empty_sections, (model, config)),
        (validate_date_format_and_values, (model, config)),
        (validate_version_consistency, (model,)),
        (validate_tbd_placeholders, (model,)),
        (validate_references, (model, template_model)),
        (validate_fonts, (model, config, font_requirements)),
        (validate_whitespace, (model,)),
        (validate_author_roles, (model,)),
        (validate_personal_references, (model,)),
        (validate_revision_date_order, (model,)),
        (validate_acronym_first_use, (model,)),
        (validate_acronym_section_presence, (model, template_model)),
        (validate_spelling, (model, known_acronyms, config))
    ]

    if template_model:
        static_validators.append((validate_document_similarity, (model, template_model, config)))
        static_validators.append((validate_table_headers_consistency, (model, template_model)))

    with ThreadPoolExecutor(max_workers=min(len(static_validators), 8)) as executor:
        futures = [executor.submit(func, *args) for func, args in static_validators]
        for future in as_completed(futures):
            try:
                findings.extend(future.result())
            except Exception as e:
                logger.error(f"Validator failed: {e}", exc_info=True)
    
    # Aggregation, Consolidation & Sorting
    findings = dedupe_findings(findings)
    
    priority_order = {"BLOCKING": 0, "HIGH": 1, "MEDIUM": 2, "LOW": 3}
    findings = sorted(findings, key=lambda f: (priority_order.get(f.priority.upper(), 2), f.rule_id, f.location_id))
    
    return findings
